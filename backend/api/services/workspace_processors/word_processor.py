"""Word (.docx) processor — python-docx, paragraphs + headings + tables."""
from __future__ import annotations

import io

from .base import make_result


def process_word(file_path: str, data: bytes) -> dict:
    warnings: list[str] = []
    errors: list[str] = []
    text_parts: list[str] = []
    tables: list[dict] = []

    try:
        from docx import Document

        doc = Document(io.BytesIO(data))
        for para in doc.paragraphs:
            t = para.text.strip()
            if not t:
                continue
            style_name = (para.style.name if para.style else "") or ""
            prefix = f"[{style_name}] " if style_name.lower().startswith("heading") else ""
            text_parts.append(f"{prefix}{t}")

        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            tables.append({"index": ti, "rows": rows})

        if not text_parts and not tables:
            warnings.append("No extractable text or tables found in this Word document.")
    except Exception as exc:
        errors.append(f"Failed to read Word document: {exc}")

    return make_result(
        file_path, "docx",
        text="\n\n".join(text_parts),
        tables=tables,
        metadata={"paragraph_count": len(text_parts), "table_count": len(tables)},
        warnings=warnings, errors=errors,
    )
