"""Word / PDF / Markdown / TXT / JSON document generation for the workspace agent.

Reuses docgen.py's generic markdown + RTL rendering helpers (title-agnostic —
no coupling to the Research Studio's InnovationRecord/ResearchOutput shape)
rather than its record-specific ``build_docx``/``build_pdf`` wrappers, which
are tied to Research Studio's cover-page and mandatory-section conventions
that don't apply to arbitrary workspace-agent reports.
"""
from __future__ import annotations

import io
import json

from api.utils.docgen import (
    _add_body_docx,
    _add_heading_docx,
    _is_arabic,
    _md_to_flowables,
    _parse_md,
    _reshape_ar,
    _rl_styles,
    _set_rtl_para,
)


def build_word_document(title: str, markdown_body: str, *, lang: str = "en") -> bytes:
    from docx import Document
    from docx.shared import Cm, Pt

    is_rtl = lang == "ar" or _is_arabic(f"{title}\n{markdown_body}")
    font_name = "Arial"

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    if title:
        heading = doc.add_heading(level=0)
        run = heading.runs[0] if heading.runs else heading.add_run()
        run.text = _reshape_ar(title) if is_rtl else title
        run.font.name = font_name
        run.font.size = Pt(22)
        if is_rtl:
            _set_rtl_para(heading)

    for block in _parse_md(markdown_body):
        if block["level"] > 0 and block["title"]:
            _add_heading_docx(doc, block["title"], block["level"], is_rtl, font_name)
        if block["body"]:
            _add_body_docx(doc, block["body"], is_rtl, font_name)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_pdf_document(title: str, markdown_body: str, *, lang: str = "en") -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    is_rtl = lang == "ar" or _is_arabic(f"{title}\n{markdown_body}")
    styles = _rl_styles(is_rtl=is_rtl)

    buf = io.BytesIO()
    doc = SimpleDocTemplate(
        buf, pagesize=A4,
        topMargin=20 * mm, bottomMargin=20 * mm, leftMargin=20 * mm, rightMargin=20 * mm,
    )

    story = []
    if title:
        story.append(Paragraph(_reshape_ar(title) if is_rtl else title, styles["h1"]))
        story.append(Spacer(1, 8))
    story.extend(_md_to_flowables(markdown_body, styles, is_rtl=is_rtl))

    doc.build(story)
    return buf.getvalue()


def build_markdown_report(title: str, markdown_body: str) -> bytes:
    content = f"# {title}\n\n{markdown_body}" if title else markdown_body
    return content.encode("utf-8")


def build_txt_report(title: str, body: str) -> bytes:
    content = f"{title}\n{'=' * max(len(title), 1)}\n\n{body}" if title else body
    return content.encode("utf-8")


def build_json_report(data: dict) -> bytes:
    return json.dumps(data, indent=2, ensure_ascii=False, default=str).encode("utf-8")
