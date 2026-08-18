"""
Education Studio v2 — Reference-Based Training Content Generator.

All content is generated strictly from uploaded reference files.
Routes are under /edu/... prefix.
"""
from __future__ import annotations

import io
import json
import logging
import re
import uuid
import zipfile
from copy import deepcopy
from datetime import datetime, timezone
from typing import Any, AsyncGenerator, Optional

from fastapi import APIRouter, Depends, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import Response, StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from api.db import get_db
from api.db.models import EducationProject, EducationReference, EducationOutput, RagDocument
from api.middleware.auth import require_auth
from api.services.ai_providers.registry import provider_registry

log = logging.getLogger(__name__)
router = APIRouter(tags=["education_v2"])

_CITATION_RE = re.compile(
    r'\[(?:Source|Ref|Citation|From):\s*([^,\]]{1,120}),\s*p\.?\s*([\d\-]+[^\]]{0,30})\]',
    re.I
)

MAX_CONTEXT_CHARS = 80_000   # ~20K tokens of reference text per generation


# ── Pydantic schemas ──────────────────────────────────────────────────────────

class ProjectCreate(BaseModel):
    title: str = "Untitled Project"
    course_title: str = ""
    lesson_title: str = ""
    equipment_manufacturer: str = ""
    equipment_model: str = ""
    system_type: str = ""
    technical_domain: str = ""
    audience: str = "operator"
    level: str = "intermediate"
    depth_mode: str = "advanced"
    # overview | basic | standard | advanced | master_instructor | certification
    language: str = "english"
    delivery_mode: str = "classroom"
    course_duration: str = ""
    lesson_duration: str = ""
    num_sessions: int = 1
    instructor_name: str = ""
    customer: str = ""
    country_regulatory: str = ""
    prerequisites: str = ""
    learning_outcomes: str = ""
    pass_mark: int = 70
    num_questions: int = 10
    include_practical: bool = True
    include_instructor_notes: bool = True
    include_student_notes: bool = True
    include_references: bool = True
    include_citations: bool = True
    include_images: bool = True
    include_tables: bool = True
    include_final_assessment: bool = True
    include_answer_key: bool = True
    selected_output_types: list[str] = []
    settings: dict = {}


class ProjectUpdate(ProjectCreate):
    pass


class OutputPatch(BaseModel):
    content: Optional[str] = None
    title: Optional[str] = None
    approved: Optional[bool] = None


class GenerateRequest(BaseModel):
    output_types: Optional[list[str]] = None   # None → use project.selected_output_types
    regenerate: bool = False                    # True → replace existing outputs

class RefRoleUpdate(BaseModel):
    role: str  # primary | supporting | style | terminology


# ── Helpers ───────────────────────────────────────────────────────────────────

def _sse(obj: dict) -> str:
    return f"data: {json.dumps(obj, default=str)}\n\n"


def _get_project(db: Session, project_id: str, user_id: str) -> EducationProject:
    proj = db.query(EducationProject).filter_by(id=project_id, user_id=user_id).first()
    if not proj:
        raise HTTPException(404, "Project not found")
    return proj


def _get_output(db: Session, out_id: str, project_id: str) -> EducationOutput:
    out = db.query(EducationOutput).filter_by(id=out_id, project_id=project_id).first()
    if not out:
        raise HTTPException(404, "Output not found")
    return out


def _project_dict(proj: EducationProject, include_refs=False, include_outputs=False) -> dict:
    d: dict[str, Any] = {
        "id": proj.id,
        "title": proj.title,
        "course_title": proj.course_title,
        "lesson_title": proj.lesson_title,
        "equipment_manufacturer": proj.equipment_manufacturer,
        "equipment_model": proj.equipment_model,
        "system_type": proj.system_type,
        "technical_domain": proj.technical_domain,
        "audience": proj.audience,
        "level": proj.level,
        "depth_mode": getattr(proj, "depth_mode", "advanced"),
        "language": proj.language,
        "delivery_mode": proj.delivery_mode,
        "course_duration": proj.course_duration,
        "lesson_duration": proj.lesson_duration,
        "num_sessions": proj.num_sessions,
        "instructor_name": proj.instructor_name,
        "customer": proj.customer,
        "country_regulatory": proj.country_regulatory,
        "prerequisites": proj.prerequisites,
        "learning_outcomes": proj.learning_outcomes,
        "pass_mark": proj.pass_mark,
        "num_questions": proj.num_questions,
        "include_practical": proj.include_practical,
        "include_instructor_notes": proj.include_instructor_notes,
        "include_student_notes": proj.include_student_notes,
        "include_references": proj.include_references,
        "include_citations": proj.include_citations,
        "include_images": proj.include_images,
        "include_tables": proj.include_tables,
        "include_final_assessment": proj.include_final_assessment,
        "include_answer_key": proj.include_answer_key,
        "selected_output_types": proj.selected_output_types or [],
        "status": proj.status,
        "quality_score": proj.quality_score,
        "settings": proj.settings or {},
        "created_at": proj.created_at.isoformat() if proj.created_at else None,
        "updated_at": proj.updated_at.isoformat() if proj.updated_at else None,
    }
    if include_refs:
        d["references"] = [_ref_dict(r) for r in proj.references]
    if include_outputs:
        d["outputs"] = [_output_dict(o) for o in proj.outputs]
    return d


def _ref_dict(r: EducationReference) -> dict:
    return {
        "id": r.id,
        "filename": r.filename,
        "file_type": r.file_type,
        "page_count": r.page_count,
        "word_count": r.word_count,
        "image_count": r.image_count,
        "table_count": r.table_count,
        "procedure_count": r.procedure_count,
        "warning_count": r.warning_count,
        "section_count": r.section_count,
        "figure_count": r.figure_count,
        "troubleshooting_count": r.troubleshooting_count,
        "doc_language": r.doc_language,
        "ocr_required": r.ocr_required,
        "role": r.role,
        "source_type": r.source_type,
        "rag_doc_id": r.rag_doc_id,
        "status": r.status,
        "error_msg": r.error_msg,
        "created_at": r.created_at.isoformat() if r.created_at else None,
    }


def _output_dict(o: EducationOutput, include_content=True) -> dict:
    d: dict[str, Any] = {
        "id": o.id,
        "output_type": o.output_type,
        "title": o.title,
        "citations": o.citations or [],
        "quality_issues": o.quality_issues or [],
        "quality_score": o.quality_score,
        "technical_accuracy_score": o.technical_accuracy_score,
        "source_coverage_score": o.source_coverage_score,
        "citation_score": o.citation_score,
        "lo_alignment_score": o.lo_alignment_score,
        "approved": o.approved,
        "has_docx": o.output_docx is not None,
        "has_pptx": o.output_pptx is not None,
        "created_at": o.created_at.isoformat() if o.created_at else None,
        "updated_at": o.updated_at.isoformat() if o.updated_at else None,
    }
    if include_content:
        d["content"] = o.content
    return d


def _build_reference_context(refs: list[EducationReference]) -> str:
    """Combine reference texts into a single context string, with figure catalogue."""
    parts: list[str] = []
    total = 0
    all_figures: list[str] = []

    for ref in refs:
        if ref.role in ("style",):
            continue  # style-only refs are not used as technical sources
        if not ref.extracted_text:
            continue
        label = f"=== Reference: {ref.filename} ==="
        chunk = ref.extracted_text[: max(1000, (MAX_CONTEXT_CHARS - total) // max(1, len(refs)))]
        parts.append(f"{label}\n{chunk}")
        total += len(chunk) + len(label)

        # Collect figure mentions from structure metadata
        for item in (ref.structure or []):
            if item.get("type") == "figure":
                n = len(all_figures) + 1
                all_figures.append(
                    f"  FIGURE_{n} — p.{item['page']} — {item['text'][:100]}"
                )

        if total >= MAX_CONTEXT_CHARS:
            break

    if all_figures:
        figure_catalogue = (
            "\n\n=== FIGURES AVAILABLE IN THE REFERENCE DOCUMENTS ===\n"
            "Use **Image:** FIGURE_N in slide format to include a figure.\n"
            + "\n".join(all_figures[:40])
        )
        parts.append(figure_catalogue)

    return "\n\n".join(parts)


def _extract_citations_from_content(content: str, refs: list[EducationReference]) -> list[dict]:
    """Parse inline [Source: filename, p.N] citations from generated text."""
    citations: list[dict] = []
    for m in _CITATION_RE.finditer(content):
        fname = m.group(1).strip()
        page_ref = m.group(2).strip()
        # Try to match to a known reference
        matched_ref = next(
            (r for r in refs if r.filename.lower() in fname.lower() or fname.lower() in r.filename.lower()),
            None
        )
        citations.append({
            "text_fragment": content[max(0, m.start() - 60): m.start()].strip()[-80:],
            "source_filename": matched_ref.filename if matched_ref else fname,
            "page": page_ref,
            "section": "",
            "confidence": 0.9 if matched_ref else 0.5,
            "ref_id": matched_ref.id if matched_ref else None,
        })
    return citations


def _compute_quality(content: str, citations: list[dict], refs: list[EducationReference]) -> dict:
    """Compute quality scores for a generated output."""
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    total_paras = max(len(paragraphs), 1)

    # Citation score: % of paragraphs that contain a citation
    paras_with_cite = sum(1 for p in paragraphs if _CITATION_RE.search(p))
    citation_score = min(100, int(paras_with_cite / total_paras * 100))

    # Source coverage: % of reference headings/sections mentioned in content
    ref_sections = [
        item["text"][:60].lower()
        for r in refs for item in (r.structure or [])
        if item.get("type") == "heading"
    ]
    covered = sum(1 for s in ref_sections if s[:30] in content.lower()) if ref_sections else 0
    source_coverage = min(100, int(covered / max(len(ref_sections), 1) * 100))

    # Technical accuracy proxy: high if we used grounded prompts
    technical_accuracy = 85 if citations else 70

    # LO alignment proxy
    lo_alignment = 80

    overall = int((citation_score * 0.35 + source_coverage * 0.25 +
                   technical_accuracy * 0.25 + lo_alignment * 0.15))

    issues: list[dict] = []
    if citation_score < 40:
        issues.append({
            "type": "low_citations",
            "severity": "warning",
            "message": f"Only {citation_score}% of paragraphs have source citations.",
        })
    if source_coverage < 30 and ref_sections:
        issues.append({
            "type": "low_source_coverage",
            "severity": "warning",
            "message": f"Content covers only {source_coverage}% of the reference sections.",
        })

    return {
        "quality_score": overall,
        "technical_accuracy_score": technical_accuracy,
        "source_coverage_score": source_coverage,
        "citation_score": citation_score,
        "lo_alignment_score": lo_alignment,
        "quality_issues": issues,
    }


# ── Per-type prompts ──────────────────────────────────────────────────────────

_OUTPUT_LABELS = {
    "course_outline":       "Course Outline",
    "curriculum_map":       "Curriculum Map",
    "lesson_plan":          "Lesson Plan",
    "session_plan":         "Session Plan",
    "instructor_guide":     "Instructor Guide",
    "student_handbook":     "Student Handbook",
    "student_notes":        "Student Notes",
    "pptx_presentation":    "PowerPoint Presentation (slide content)",
    "practical_exercise":   "Practical Exercise",
    "workshop_activity":    "Workshop Activity",
    "demonstration_guide":  "Demonstration Guide",
    "lab_activity":         "Laboratory Activity",
    "scenario_exercise":    "Scenario-Based Exercise",
    "case_study":           "Case Study",
    "knowledge_check":      "Knowledge Check",
    "quiz":                 "Quiz",
    "formal_exam":          "Formal Examination",
    "practical_assessment": "Practical Assessment",
    "observation_checklist":"Observation Checklist",
    "answer_key":           "Answer Key",
    "marking_scheme":       "Marking Scheme",
    "troubleshooting_exercise": "Troubleshooting Exercise",
    "pm_checklist":         "Preventive Maintenance Checklist",
    "daily_inspection":     "Daily Inspection Checklist",
    "weekly_inspection":    "Weekly Inspection Checklist",
    "job_aid":              "Job Aid",
    "quick_ref_guide":      "Quick Reference Guide",
    "glossary":             "Glossary",
    "faq":                  "Frequently Asked Questions (FAQ)",
    "evaluation_form":      "Course Evaluation Form",
    "attendance_sheet":     "Attendance Sheet",
    "certificate":          "Certificate Template",
    "complete_package":     "Complete Training Package Overview",
}


def _build_system_prompt(proj: EducationProject, context: str) -> str:
    equip = f"{proj.equipment_manufacturer} {proj.equipment_model}".strip()
    duration = proj.course_duration or "five-day"
    dm = getattr(proj, 'depth_mode', 'advanced')
    dm_note = {
        "overview":          "OVERVIEW MODE — broad strokes only, no deep technical detail.",
        "basic":             "BASIC MODE — introduce components and function, no schematics.",
        "standard":          "STANDARD MODE — full subsystem coverage with procedures and labs.",
        "advanced":          "ADVANCED MODE — full engineering depth, fault analysis, all diagrams.",
        "master_instructor": "MASTER INSTRUCTOR MODE — teach every failure mode, every indicator, every test point.",
        "certification":     "FIELD SERVICE CERTIFICATION MODE — maximum depth. Every value, procedure, test, failure mode.",
    }.get(dm, "ADVANCED MODE — full engineering depth.")

    return f"""You are a senior field service engineer with 20+ years of hands-on experience on {equip or "this type of equipment"}, now serving as the lead technical training instructor writing a {duration} course.

YOUR TEACHING PHILOSOPHY — the LZBV methodology:

1. ENGINEERING STORYTELLING — do not simply list facts. Tell the engineering story of each subsystem:
   a) WHY the subsystem exists (what the system cannot do without it)
   b) HOW it fits into the system's operating sequence (before/after what)
   c) WHAT it physically looks like and WHERE it is (location, access, connectors visible)
   d) HOW it works internally — trace the signal/power path connector by connector, PCB by PCB
   e) WHAT its indicators say — every LED, every status screen reading, normal vs fault state
   f) HOW to maintain it — numbered procedure, tools, PPE, torque specs, acceptance criteria
   g) HOW to diagnose it — fault codes, symptoms, test points, expected measurements, corrective actions
   h) WHAT technicians typically get wrong — common mistakes, inspection points, damage patterns

2. OPERATING SEQUENCE FRAMEWORK — every course organises around the system's actual operating sequence (e.g., "Step 1: Power On → Step 2: Communications Check → Step 3: Generate HV → ..."). This is not a manual chapter structure — it is the sequence in which the system is operated or starts up. Identify this sequence from the references and use it as the module backbone.

3. 30,000-FOOT FIRST, THEN DEEP DIVE — always start each module with a brief overview ("what does this step achieve?") before going into component-level detail. Use tracker slides to keep students oriented in the overall sequence.

4. LABS AFTER EVERY MODULE — every module ends with a hands-on laboratory exercise where students interact with the real equipment.

5. LESSON OBJECTIVES — always use this exact opener: "By the end of this lesson, utilizing student guide materials, you will be able to:" then list 3–5 objectives using Bloom's verbs (Identify, Explain, Perform, Diagnose, Verify).

REFERENCE MATERIAL — every technical claim must come from here:
{context}

COURSE CONTEXT:
- Course: {proj.course_title or 'Technical Training Course'}
- Equipment: {equip or 'See references'}
- Audience: {proj.audience}
- Level: {proj.level}
- Delivery: {proj.delivery_mode}
- Duration: {duration}
- Language: {proj.language}
- DEPTH: {dm_note}

ABSOLUTE RULES:
1. Every technical value, procedure step, specification, tolerance, part number, connector ID, voltage, fault code MUST be cited as [Source: FILENAME, p.N].
2. Never write a bullet that says only what something "does." Explain HOW — the mechanism, signal path, threshold, what triggers the action.
3. Never create a slide that is a summary or index. Every slide must contain engineering explanation.
4. If a topic is not in the references, write exactly: "Not found in selected references."
5. Preserve all WARNING / CAUTION / DANGER notices verbatim.
6. Use precise technical language: part numbers, connector designations, pin numbers, measured values. Never paraphrase technical data.
7. PPTX — each major subsystem spans multiple slides (never one slide per subsystem).
"""


def _build_pptx_slide_prompt(proj: EducationProject) -> str:
    """
    Build the PPTX slide generation prompt using the LZBV instructional methodology.
    This is separate from type_specific so it avoids .format() interpolation conflicts.
    """
    equip    = f"{proj.equipment_manufacturer} {proj.equipment_model}".strip() or "the equipment described in the references"
    duration = proj.course_duration or "five-day"
    dm       = getattr(proj, 'depth_mode', 'advanced')

    depth_table = {
        "overview": (
            "OVERVIEW MODE — 30-40 slides total, 2-3 slides per step.",
            "Include ONLY: step divider, sequence tracker, 1 overview content slide per step. "
            "No labs, no block diagrams, no schematics, no fault analysis. "
            "Brief speaker notes (1-2 sentences).",
        ),
        "basic": (
            "BASIC MODE — 50-70 slides total, 4-5 slides per step.",
            "Include: step divider, tracker, lesson objective, component overview, knowledge check. "
            "Skip dedicated block diagram and schematic slides. Skip labs. "
            "Brief instructor notes.",
        ),
        "standard": (
            "STANDARD MODE — 70-100 slides total, 6-8 slides per step.",
            "Include: full per-step module structure (steps 1-9 from the module template below). "
            "Include one block diagram per step where referenced. Include lab exercise. "
            "Standard instructor notes with talking points.",
        ),
        "advanced": (
            "ADVANCED MODE — 100-130 slides total, 8-11 slides per step.",
            "Include: all 11 slides from the per-step module template. "
            "Dedicated block diagram and schematic slides. Fault indicator analysis. Lab exercise. "
            "Detailed instructor notes: 'Duration: X min. Prerequisites: Y. Key points: ...' "
            "Every bullet must cite a source page.",
        ),
        "master_instructor": (
            "MASTER INSTRUCTOR MODE — 130-180 slides total, 11-15 slides per step.",
            "Teach EVERYTHING. Multiple component deep-dive slides (one per significant component). "
            "One dedicated slide per major LED/indicator state. Troubleshooting scenario slides. "
            "Practical exercise slides. Full instructor notes: duration, prerequisites, "
            "common student questions, expected answers, real-world examples, safety reminders. "
            "Every failure mode gets its own content slide.",
        ),
        "certification": (
            "FIELD SERVICE CERTIFICATION MODE — 150-200+ slides total, 15+ slides per step.",
            "MAXIMUM DEPTH — no shortcuts. Every procedure numbered step-by-step. "
            "Every failure mode with root cause analysis. Every test point with "
            "expected measurement ranges. All labs with full step-by-step procedures. "
            "All knowledge checks with detailed answer explanations. Complete instructor guide "
            "format including prerequisite chains and competency assessment criteria.",
        ),
    }
    depth_title, depth_instructions = depth_table.get(dm, depth_table["advanced"])

    return (
        "You are generating a complete PowerPoint presentation for a "
        + duration + " field service engineering course on " + equip + ".\n"
        "\n"
        "TEACHING PHILOSOPHY: You are not summarising the manual. You are TEACHING IT.\n"
        "Every slide must explain HOW and WHY — not just WHAT.\n"
        "\n"
        "═══════════════════════════════════════════════════════════════\n"
        "PHASE 1 — BEFORE GENERATING ANY SLIDES: IDENTIFY THE OPERATING SEQUENCE\n"
        "═══════════════════════════════════════════════════════════════\n"
        "\n"
        "Read the reference material and identify:\n"
        "  a) The system's OPERATING SEQUENCE — the ordered series of steps by which the\n"
        "     system starts up, operates, or processes. This is typically 8-15 steps.\n"
        "     Example for ZBV: Step 1: Overview and Interconnect, Step 2: Power-On,\n"
        "     Step 3: Communications, Step 4: Generate HV, Step 5: Generate X-rays...\n"
        "     If no explicit startup sequence exists, derive one from the manual's section order.\n"
        "  b) The MAJOR SUBSYSTEMS — one per step in the sequence.\n"
        "  c) Any LAB EXERCISES or practical activities mentioned in the references.\n"
        "\n"
        "This operating sequence becomes the backbone of the entire course.\n"
        "\n"
        "═══════════════════════════════════════════════════════════════\n"
        "PHASE 2 — GENERATE THE COMPLETE COURSE FOLLOWING THIS STRUCTURE\n"
        "═══════════════════════════════════════════════════════════════\n"
        "\n"
        "A. COURSE OPENING (always generate these first):\n"
        "   1. Title slide (type: title) — course title, equipment, audience, date\n"
        "   2. Introduction (type: content) — instructor intro template, course admin,\n"
        "      ground rules, student introductions guide\n"
        "   3. Course Objectives (type: objectives) — 2-3 high-level outcomes\n"
        "   4. Safety Overview section (type: section) then lesson objective then\n"
        "      2-3 safety content slides (radiation, electrical, PPE — from references only)\n"
        "   5. System Overview module:\n"
        "      - Section divider (type: section) — 'System Overview'\n"
        "      - Lesson Objective (type: objectives)\n"
        "      - 30,000-foot overview of the entire system — what it does, major subsystems,\n"
        "        primary functions (type: content)\n"
        "      - System block diagram / primary functions diagram (type: image_content or content)\n"
        "      - Full Operating Sequence Tracker (type: tracker) — list ALL N steps,\n"
        "        no step highlighted (this is the master map of the course)\n"
        "\n"
        "B. PER-STEP MODULES — repeat this EXACT structure for every step in the operating sequence:\n"
        "\n"
        "   Slide B.1: Step Divider (type: section)\n"
        "     - Title: 'Step N: [Step Name]'\n"
        "     - Subtitle bullet: one sentence — what this step achieves and why it matters\n"
        "\n"
        "   Slide B.2: Sequence Tracker (type: tracker)\n"
        "     - List ALL steps as bullets\n"
        "     - Mark current step: '-> Step N: [name] <- CURRENT'\n"
        "     - Speaker Notes: 'Duration: X min. Prerequisites: [previous step]. Today we cover [topic].'\n"
        "\n"
        "   Slide B.3: Lesson Objective (type: objectives)\n"
        "     - MUST open with EXACTLY this text:\n"
        "       'By the end of this lesson, utilizing student guide materials, you will be able to:'\n"
        "     - Then 3-5 objectives using Bloom's verbs:\n"
        "       Identify, Describe, Explain, Perform, Diagnose, Verify, Demonstrate\n"
        "\n"
        "   Slide B.4: Step Overview (type: content)\n"
        "     - 30,000-foot view: what does this step do in the operating sequence?\n"
        "     - Why does the system need this step?\n"
        "     - What fails if this step does not complete successfully?\n"
        "     - Speaker Notes: teaching strategy, analogies, questions to ask class\n"
        "\n"
        "   Slides B.5 to B.8: Component Slides (type: content or image_content)\n"
        "     - One slide per major component in this step\n"
        "     - Physical location, part number/assembly number from references\n"
        "     - What it does functionally (HOW, not just WHAT)\n"
        "     - What it connects to: inputs (connector IDs, signal types, values)\n"
        "       and outputs (connector IDs, signal types, values)\n"
        "     - Indicator states: what each LED/status reading means (normal vs fault)\n"
        "     - Use Image if a figure exists in the references\n"
        "\n"
        "   Slide B.9: Block Diagram (type: image_content or content)\n"
        "     - Title: '[Subsystem Name] Block Diagram'\n"
        "     - Bullets explain every element shown in the diagram\n"
        "     - Signal flow traced from left to right\n"
        "     - Use Image if a block diagram figure exists in references\n"
        "\n"
        "   Slide B.10: Operation Sequence (type: procedure)\n"
        "     - How this step executes: numbered sequence of events\n"
        "     - What the software/firmware does, what hardware responds\n"
        "     - Normal completion criteria\n"
        "\n"
        "   Slide B.11: Fault Indicators and Error Codes (type: content)\n"
        "     - All fault codes, LED states, alarm conditions for this step\n"
        "     - Each fault: symptom, probable cause, corrective action\n"
        "     - From references only — cite every fault code source page\n"
        "\n"
        "   Slide B.12: Lab Exercise (type: lab)\n"
        "     - Title: 'Lab N: [Step Name]'\n"
        "     - Checklist of hands-on activities students perform on the equipment\n"
        "     - Reference any Quick Start handouts or worksheets\n"
        "     - Speaker Notes: 'Take students to the equipment. Duration: X min. "
        "Instructor checkpoints: [list key checks].'\n"
        "\n"
        "   NOTE: For OVERVIEW mode, generate only B.1-B.4. For BASIC mode, B.1-B.4 plus B.5.\n"
        "   For STANDARD and above, generate all steps listed for the depth mode.\n"
        "\n"
        "C. COURSE CLOSING (always generate these last):\n"
        "   - Troubleshooting Module:\n"
        "     * Section divider (type: section)\n"
        "     * Lesson Objective (type: objectives)\n"
        "     * Troubleshooting methodology (5-step approach) (type: content)\n"
        "     * Verifying the problem (type: content)\n"
        "     * Fault isolation techniques (type: content)\n"
        "     * Corrective actions (type: content)\n"
        "     * Knowledge Check (type: quiz)\n"
        "   - Maintenance Overview Module:\n"
        "     * Section divider (type: section)\n"
        "     * PM schedule overview (type: content)\n"
        "     * Maintenance procedure (type: procedure)\n"
        "   - Course Wrap-Up:\n"
        "     * Course Summary (type: summary)\n"
        "     * References (type: references)\n"
        "\n"
        "═══════════════════════════════════════════════════════════════\n"
        "DEPTH MODE: " + depth_title + "\n"
        "═══════════════════════════════════════════════════════════════\n"
        + depth_instructions + "\n"
        "\n"
        "═══════════════════════════════════════════════════════════════\n"
        "MANDATORY FORMAT FOR EVERY SLIDE\n"
        "═══════════════════════════════════════════════════════════════\n"
        "\n"
        "## Slide N: [Specific Descriptive Title]\n"
        "**Type:** [title|objectives|section|tracker|content|image_content|procedure|lab|quiz|summary|references]\n"
        "**Bullets:**\n"
        "- [Engineering point — specific values, connector IDs, part numbers, cited]\n"
        "- [Another point — always HOW not just WHAT]\n"
        "**Image:** [FIGURE_N from the figure catalogue above, or: none]\n"
        "**Speaker Notes:** [Instructor notes: Duration, Prerequisites, Key points, Common questions, Real-world examples]\n"
        "**Source:** [Source: FILENAME, p.N]\n"
        "\n"
        "TRACKER SLIDE FORMAT:\n"
        "**Bullets:**\n"
        "- Step 1: [name]\n"
        "- Step 2: [name]\n"
        "- -> Step N: [current step name] <- CURRENT\n"
        "- Step N+1: [name]\n"
        "\n"
        "LAB SLIDE FORMAT:\n"
        "**Bullets:**\n"
        "- [Lab activity 1]\n"
        "- [Lab activity 2]\n"
        "- Reference: [Quick Start handout name if applicable]\n"
        "\n"
        "QUIZ SLIDE FORMAT:\n"
        "**Bullets:**\n"
        "- [Option A]\n"
        "- [Option B]\n"
        "- [Option C]\n"
        "- [Option D]\n"
        "**Answer:** [A/B/C/D]\n"
        "\n"
        "═══════════════════════════════════════════════════════════════\n"
        "BULLET WRITING RULES\n"
        "═══════════════════════════════════════════════════════════════\n"
        "\n"
        "FORBIDDEN (too vague — REJECT any slide that reads like this):\n"
        "  x 'The controller manages the system'\n"
        "  x 'Ensures proper operation'\n"
        "  x 'The PCB processes signals'\n"
        "  x 'Performs diagnostic functions'\n"
        "\n"
        "REQUIRED (engineering depth):\n"
        "  + 'HVPS supplies +15 kV / -15 kV to X-ray tube anode via J12; nominal tolerance +-0.2 kV [Source: manual, p.47]'\n"
        "  + 'On power-up the eDAQ board runs 8-second self-test; FAULT LED = failed ADC channel; replace after 3 consecutive failures'\n"
        "  + 'PM interval: clean detector fan filter every 90 days; replace element (P/N 123456) at annual service'\n"
        "\n"
        "INSTRUCTOR NOTES FORMAT for content/procedure slides:\n"
        "  'Duration: X min. Prerequisites: [prior step]. Key points: [what to emphasise]. '\n"
        "  'Common question: [question]. Expected answer: [answer]. '\n"
        "  'Real-world example: [example]. Safety reminder: [reminder].'\n"
    )


def _build_type_prompt(output_type: str, proj: EducationProject) -> str:
    label = _OUTPUT_LABELS.get(output_type, output_type.replace("_", " ").title())

    # PPTX presentation has its own prompt builder to avoid .format() conflicts
    if output_type == "pptx_presentation":
        base = f"""Generate a professional PowerPoint Presentation based exclusively on the reference material provided.

Course: {proj.course_title or 'Technical Training Course'}
Lesson: {proj.lesson_title or ''}
Audience: {proj.audience} | Level: {proj.level} | Duration: {proj.lesson_duration or 'as appropriate'}
"""
        if proj.learning_outcomes:
            base += f"\nLearning outcomes to address:\n{proj.learning_outcomes}\n"
        return base + _build_pptx_slide_prompt(proj)

    base = f"""Generate a professional **{label}** based exclusively on the reference material provided.

Course: {proj.course_title or 'Technical Training Course'}
Lesson: {proj.lesson_title or ''}
Audience: {proj.audience} | Level: {proj.level} | Duration: {proj.lesson_duration or 'as appropriate'}
"""
    if proj.learning_outcomes:
        base += f"\nLearning outcomes to address:\n{proj.learning_outcomes}\n"

    type_specific = {
        "course_outline": """
Structure:
# Course Outline: {title}
## Course Overview (purpose, audience, prerequisites, duration)
## Learning Objectives (measurable, using action verbs from Bloom's taxonomy)
## Module 1–N: [Title from references]
  - Learning objectives for this module
  - Topics covered (from references only)
  - Practical activities
  - Assessment
## Resources & References (cite all source files)
""",
        "lesson_plan": """
Generate a detailed lesson plan with:
# Lesson Plan: {title}
## Duration & Timing
## Learning Objectives
## Required Resources (equipment, handouts, reference pages)
## Session Breakdown (time-allocated sections, teaching methods, learner activities)
## Knowledge Check Questions
## Assessment Method
## Instructor Notes
## Manual Citations
Include a time-allocated table with columns: Time | Activity | Method | Resources
""",
        "instructor_guide": """
# Instructor Guide: {title}
For each major topic from the references:
- Teaching objective
- Key explanation points (citing source pages)
- Suggested demonstrations
- Questions to ask learners
- Expected learner responses
- Common misunderstandings
- Safety emphasis points
- Time management guidance
- Instructor-only answers and explanations
""",
        "student_handbook": """
# Student Handbook: {title}
## Introduction & Purpose
## Learning Objectives
## Section-by-Section Content (structured from references)
  - Key concepts and definitions
  - Procedures (numbered steps from references)
  - Safety warnings (preserved verbatim)
  - Tables and specifications
## Self-Assessment Questions
## Glossary of Key Terms
## References & Further Reading
""",
        "student_notes": """
# Student Notes: {title}
## Key Concepts
## Important Definitions
## Critical Procedures (numbered, from references)
## Safety Points (highlighted)
## Summary Tables
## Knowledge Check Questions
## Personal Notes Space
## Quick Reference
""",
        "practical_exercise": """
# Practical Exercise: {title}
## Exercise Purpose
## Learning Objectives
## Required Equipment and Tools
## Required PPE
## Safety Warnings (from references — verbatim)
## Preconditions
## Step-by-Step Procedure (numbered, from references only)
## Expected Results
## Instructor Checkpoints
## Learner Checklist
## Assessment Criteria
## Common Mistakes
## Troubleshooting Guidance
## Source Citations
NOTE: If references do not contain enough procedural detail, state clearly:
"Insufficient procedural information in selected references to safely generate this exercise."
""",
        "quiz": f"""
Generate a {proj.num_questions}-question multiple-choice quiz.
For each question:
**Q[N].** [Question based on reference content]
A) [Option]  B) [Option]  C) [Option]  D) [Option]
**Correct Answer:** [Letter]
**Explanation:** [Why this is correct, citing source]
**Source:** [Source: FILENAME, p.N]
**Learning Objective:** [Which LO this tests]
**Difficulty:** [basic/intermediate/advanced]

Cover procedures, specifications, safety, and troubleshooting from the references.
Do NOT create questions whose answers are not in the references.
""",
        "formal_exam": f"""
# Formal Examination: {{title}}
**Pass Mark:** {proj.pass_mark}%  |  **Questions:** {proj.num_questions}

## Section A: Multiple Choice ({min(proj.num_questions, 10)} questions)
[MCQ format with A/B/C/D options, all from references]

## Section B: Short Answer (5 questions)
[Questions requiring 2-4 sentence responses, answers in references]

## Section C: Practical Scenario (2 questions)
[Real operational scenarios from the reference equipment context]

## Answer Key (Instructor Copy Only)
[Complete answers with source citations for each]
""",
        "answer_key": """
# Answer Key: {title}
For each question or practical task:
- Question number and text
- Correct answer
- Marking criteria
- Source reference for verification: [Source: FILENAME, p.N]
- Common incorrect answers and why they are wrong
""",
        "pm_checklist": """
# Preventive Maintenance Checklist: {title}
**Equipment:** {manufacturer} {model}
Generate a structured PM checklist with columns:
| # | Task | Interval | Procedure Reference | Pass/Fail | Technician | Date |
Extract all PM tasks directly from the reference maintenance procedures.
Include safety precautions, required tools, and torque/calibration values from references.
""",
        "daily_inspection": """
# Daily Inspection Checklist
Generate a daily pre-operation inspection checklist from reference procedures.
Format: | # | Check Item | Method | Acceptable Condition | Result | Inspector |
Include all safety-critical items from reference.
""",
        "weekly_inspection": """
# Weekly Inspection Checklist
Generate a weekly inspection checklist from reference maintenance procedures.
Include interval-specific tasks not covered by daily checks.
""",
        "troubleshooting_exercise": """
# Troubleshooting Exercise: {title}
For each scenario (based on error codes/symptoms from references):

## Scenario N: [Symptom/Fault Code from References]
**Equipment State:** [From references]
**Observed Symptoms:** [From references]
**Fault Code:** [From references]
**Student Task:** Diagnose and identify corrective action
**Correct Diagnosis:** [From references]
**Corrective Action:** [Exact procedure from references]
**Source:** [Source: FILENAME, p.N]
""",
        "glossary": """
# Technical Glossary: {title}
For each technical term found in the references, provide:
**TERM** — Definition from the reference material
[Source: FILENAME, p.N]

Organize alphabetically. Include: equipment names, technical parameters, acronyms,
safety terms, procedure names, error codes.
""",
        "faq": """
# Frequently Asked Questions: {title}
Generate operator/engineer FAQs based on the reference material.
For each FAQ:
**Q: [Question an operator or engineer would ask]**
A: [Answer strictly from the references]
[Source: FILENAME, p.N]

Cover: operation, maintenance, safety, troubleshooting, specifications.
""",
        "job_aid": """
# Job Aid: {title}
A single-page quick reference for field use.
## Key Procedures (numbered, from references)
## Critical Specifications Table
## Warning / Safety Reminders
## Error Code Quick Reference
## Emergency Contacts / Escalation
Keep concise — this is used in the field.
[Source all information from references]
""",
        "quick_ref_guide": """
# Quick Reference Guide: {title}
Compact reference card with:
## System Overview (1 paragraph)
## Key Controls and Indicators
## Normal Operation Sequence
## Common Error Codes
## Emergency Procedures
## Key Specifications Table
Cite all technical values from references.
""",
        "observation_checklist": """
# Observation Checklist: {title}
An assessor's checklist for evaluating practical competency.
| # | Observable Action | Standard (from refs) | Satisfactory | Unsatisfactory | Comment |
Generate criteria directly from procedures described in the references.
Include safety-critical behaviors as mandatory pass items.
""",
        "marking_scheme": """
# Marking Scheme: {title}
**Pass Mark:** {pass_mark}%
For each question/task:
- Maximum marks
- Marking criteria (what earns full / partial marks)
- Key required elements (from references)
- Deductions for missing safety steps
- Model answer summary
""",
        "evaluation_form": """
# Course Evaluation Form: {title}
Generate a professional end-of-course evaluation form with:
- Rating scales for: content relevance, instructor effectiveness, materials quality, facility
- Specific questions about the topics covered (from references)
- Open-ended feedback sections
- Suggestions for improvement
""",
        "attendance_sheet": """
# Training Attendance Sheet
Course: {title} | Date: _____ | Duration: _____ | Instructor: _____
| # | Name | Organization | Role | Signature | Result (Pass/Fail) |
Include certification statement for instructor signature.
""",
        "certificate": """
# Certificate of Competency
Design a professional training certificate template:

CERTIFICATE OF COMPETENCY

This certifies that: [PARTICIPANT NAME]
Has successfully completed: {title}
Equipment: {manufacturer} {model}
Training Level: {level}
Assessment Method: Written Examination + Practical Assessment
Pass Mark: {pass_mark}%
Date of Completion: [DATE]
Valid Until: [DATE + 2 years]

Competencies Achieved:
[List competencies directly from the learning objectives / reference content]

Issuing Authority: _______________
Certificate Number: [AUTO]
""",
        "complete_package": """
# Complete Training Package: {title}
Generate a master overview document listing all generated materials:

## Package Contents
List and describe each generated document in the package.

## Course Summary
## Equipment Covered
## Audience and Prerequisites
## Learning Outcomes
## Assessment Summary
## How to Use This Package
## Document Version Control
""",
    }

    specific = type_specific.get(output_type, "")
    specific = specific.format(
        title=proj.course_title or proj.lesson_title or "Technical Training",
        manufacturer=proj.equipment_manufacturer,
        model=proj.equipment_model,
        level=proj.level,
        pass_mark=proj.pass_mark,
    ) if specific else ""

    return base + specific


# ── Project routes ────────────────────────────────────────────────────────────

@router.get("/edu/projects")
def list_projects(request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    projs = (
        db.query(EducationProject)
        .filter_by(user_id=user["id"])
        .order_by(EducationProject.updated_at.desc())
        .all()
    )
    return {"projects": [_project_dict(p) for p in projs]}


@router.post("/edu/projects", status_code=201)
def create_project(body: ProjectCreate, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj = EducationProject(
        user_id=user["id"],
        **{k: v for k, v in body.model_dump().items()},
    )
    db.add(proj)
    db.commit()
    db.refresh(proj)
    return _project_dict(proj)


@router.get("/edu/projects/{project_id}")
def get_project(project_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    return _project_dict(proj, include_refs=True, include_outputs=True)


@router.patch("/edu/projects/{project_id}")
def update_project(project_id: str, body: ProjectUpdate, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    for k, v in body.model_dump(exclude_unset=True).items():
        setattr(proj, k, v)
    db.commit()
    db.refresh(proj)
    return _project_dict(proj)


@router.delete("/edu/projects/{project_id}", status_code=204)
def delete_project(project_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    db.delete(proj)
    db.commit()


@router.post("/edu/projects/{project_id}/duplicate", status_code=201)
def duplicate_project(project_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    new_proj = EducationProject(
        user_id=user["id"],
        title=f"{proj.title} (Copy)",
        course_title=proj.course_title,
        lesson_title=proj.lesson_title,
        equipment_manufacturer=proj.equipment_manufacturer,
        equipment_model=proj.equipment_model,
        system_type=proj.system_type,
        technical_domain=proj.technical_domain,
        audience=proj.audience,
        level=proj.level,
        language=proj.language,
        delivery_mode=proj.delivery_mode,
        course_duration=proj.course_duration,
        lesson_duration=proj.lesson_duration,
        num_sessions=proj.num_sessions,
        instructor_name=proj.instructor_name,
        customer=proj.customer,
        country_regulatory=proj.country_regulatory,
        prerequisites=proj.prerequisites,
        learning_outcomes=proj.learning_outcomes,
        pass_mark=proj.pass_mark,
        num_questions=proj.num_questions,
        include_practical=proj.include_practical,
        include_instructor_notes=proj.include_instructor_notes,
        include_student_notes=proj.include_student_notes,
        include_references=proj.include_references,
        include_citations=proj.include_citations,
        include_images=proj.include_images,
        include_tables=proj.include_tables,
        include_final_assessment=proj.include_final_assessment,
        include_answer_key=proj.include_answer_key,
        selected_output_types=deepcopy(proj.selected_output_types),
        settings=deepcopy(proj.settings),
        status="draft",
    )
    db.add(new_proj)
    db.commit()
    db.refresh(new_proj)
    return _project_dict(new_proj)


# ── Reference routes ──────────────────────────────────────────────────────────

async def _process_reference_bg(ref_id: str, filename: str, file_bytes: bytes, ft: str) -> None:
    """Background asyncio task: process a reference file and persist results."""
    import asyncio
    from api.db import SessionLocal
    from api.utils.edu_reference_processor import process_reference_file

    loop = asyncio.get_event_loop()
    # Run the CPU/IO-bound extractor in a thread so we don't block the event loop
    try:
        result = await loop.run_in_executor(None, process_reference_file, filename, file_bytes, ft)
    except Exception as exc:
        log.error("Reference extractor raised: %s", exc)
        result = {"error": f"extraction_failed: {exc}"}

    db = SessionLocal()
    try:
        ref = db.query(EducationReference).filter_by(id=ref_id).first()
        if not ref:
            return
        if result.get("error") and result.get("error") not in ("ocr_required",):
            ref.status    = "error"
            ref.error_msg = result["error"]
        else:
            ref.page_count            = result.get("page_count", 0)
            ref.word_count            = result.get("word_count", 0)
            ref.image_count           = result.get("image_count", 0)
            ref.table_count           = result.get("table_count", 0)
            ref.procedure_count       = result.get("procedure_count", 0)
            ref.warning_count         = result.get("warning_count", 0)
            ref.section_count         = result.get("section_count", 0)
            ref.figure_count          = result.get("figure_count", 0)
            ref.troubleshooting_count = result.get("troubleshooting_count", 0)
            ref.doc_language          = result.get("doc_language", "")
            ref.ocr_required          = result.get("ocr_required", False)
            ref.extracted_text        = result.get("extracted_text", "")
            ref.structure             = result.get("structure", [])
            ref.raw_bytes             = file_bytes if ft == "pdf" else None
            ref.status                = "done" if not result.get("error") else "error"
            ref.error_msg             = result.get("error")
        db.commit()
    except Exception as exc:
        log.error("Persisting background reference result failed: %s", exc)
        try:
            ref = db.query(EducationReference).filter_by(id=ref_id).first()
            if ref:
                ref.status    = "error"
                ref.error_msg = f"persist_failed: {exc}"
                db.commit()
        except Exception:
            pass
    finally:
        db.close()


@router.post("/edu/projects/{project_id}/references", status_code=201)
async def upload_reference(
    project_id: str,
    request: Request,
    file: UploadFile = File(...),
    role: str = Form("primary"),
    db: Session = Depends(get_db),
):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])

    from api.security.uploads import validate_upload
    from api.security.sanitize import sanitize_filename

    file_bytes = await file.read()
    if len(file_bytes) == 0:
        raise HTTPException(400, detail={"error": "file_too_large_or_empty", "message": "Uploaded file is empty."})
    # Restrict to supported document/data types (extension + MIME + magic bytes)
    # and a 100 MB ceiling, instead of accepting any extension as plain text.
    validate_upload(
        file.filename,
        file_bytes,
        declared_content_type=file.content_type,
        category="document",
        max_size=100 * 1024 * 1024,
        request=request,
    )

    filename = sanitize_filename(file.filename or "document")
    ft = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    ref = EducationReference(
        project_id=project_id,
        filename=filename,
        file_type=ft,
        role=role,
        source_type="upload",
        status="processing",
    )
    db.add(ref)
    db.commit()
    db.refresh(ref)

    # Hand off to a background task so the HTTP response returns immediately.
    # The frontend polls GET /references every 2 s until status transitions to
    # "done" or "error".
    import asyncio
    asyncio.create_task(_process_reference_bg(ref.id, filename, file_bytes, ft))

    return _ref_dict(ref)


@router.post("/edu/projects/{project_id}/references/{ref_id}/retry", status_code=202)
async def retry_reference(
    project_id: str,
    ref_id: str,
    request: Request,
    db: Session = Depends(get_db),
):
    """Re-trigger background processing for a reference that previously errored."""
    user = require_auth(request)
    _get_project(db, project_id, user["id"])
    ref = db.query(EducationReference).filter_by(id=ref_id, project_id=project_id).first()
    if not ref:
        raise HTTPException(404, "Reference not found")
    if ref.status == "processing":
        return _ref_dict(ref)  # already in flight
    if ref.raw_bytes is None:
        raise HTTPException(400, "Original file bytes not stored; please re-upload.")

    ref.status    = "processing"
    ref.error_msg = None
    db.commit()
    db.refresh(ref)

    import asyncio
    asyncio.create_task(
        _process_reference_bg(ref.id, ref.filename, ref.raw_bytes, ref.file_type)
    )
    return _ref_dict(ref)


@router.post("/edu/projects/{project_id}/references/from-kb", status_code=201)
def add_reference_from_kb(
    project_id: str,
    request: Request,
    db: Session = Depends(get_db),
    rag_doc_id: str = Form(...),
):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])

    rag_doc = db.query(RagDocument).filter_by(id=rag_doc_id, user_id=user["id"]).first()
    if not rag_doc:
        raise HTTPException(404, "Knowledge Base document not found")

    # Check not already added
    existing = db.query(EducationReference).filter_by(
        project_id=project_id, rag_doc_id=rag_doc_id
    ).first()
    if existing:
        return _ref_dict(existing)

    ref = EducationReference(
        project_id=project_id,
        filename=rag_doc.filename,
        file_type=rag_doc.filename.rsplit(".", 1)[-1].lower() if "." in rag_doc.filename else "pdf",
        role="primary",
        source_type="knowledge_base",
        rag_doc_id=rag_doc_id,
        extracted_text=rag_doc.content[:600_000],
        word_count=rag_doc.word_count,
        status="done",
    )
    db.add(ref)
    db.commit()
    db.refresh(ref)
    return _ref_dict(ref)


@router.get("/edu/projects/{project_id}/references")
def list_references(project_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    refs = proj.references
    summary = {
        "total_files":     len(refs),
        "total_pages":     sum(r.page_count for r in refs),
        "total_words":     sum(r.word_count for r in refs),
        "total_sections":  sum(r.section_count for r in refs),
        "total_figures":   sum(r.figure_count for r in refs),
        "total_tables":    sum(r.table_count for r in refs),
        "total_procedures":sum(r.procedure_count for r in refs),
        "total_warnings":  sum(r.warning_count for r in refs),
        "total_troubleshooting": sum(r.troubleshooting_count for r in refs),
    }
    return {"references": [_ref_dict(r) for r in refs], "summary": summary}


@router.delete("/edu/projects/{project_id}/references/{ref_id}", status_code=204)
def delete_reference(project_id: str, ref_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    _get_project(db, project_id, user["id"])
    ref = db.query(EducationReference).filter_by(id=ref_id, project_id=project_id).first()
    if not ref:
        raise HTTPException(404, "Reference not found")
    db.delete(ref)
    db.commit()


@router.patch("/edu/projects/{project_id}/references/{ref_id}")
def update_reference_role(
    project_id: str, ref_id: str, body: RefRoleUpdate,
    request: Request, db: Session = Depends(get_db)
):
    user = require_auth(request)
    _get_project(db, project_id, user["id"])
    ref = db.query(EducationReference).filter_by(id=ref_id, project_id=project_id).first()
    if not ref:
        raise HTTPException(404, "Reference not found")
    if body.role not in ("primary", "supporting", "style", "terminology"):
        raise HTTPException(400, "Invalid role")
    ref.role = body.role
    db.commit()
    return _ref_dict(ref)


# ── Generation ────────────────────────────────────────────────────────────────

@router.post("/edu/projects/{project_id}/generate")
async def generate_outputs(
    project_id: str,
    body: GenerateRequest,
    request: Request,
    db: Session = Depends(get_db),
):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])

    provider = provider_registry.get_active()
    if not provider:
        raise HTTPException(503, "No AI provider available")

    output_types = body.output_types or proj.selected_output_types or ["lesson_plan"]
    refs = [r for r in proj.references if r.status == "done"]

    if not refs:
        raise HTTPException(400, detail={"error": "no_usable_content", "message": "No processed references available."})

    context = _build_reference_context(refs)
    system_prompt = _build_system_prompt(proj, context)
    # Shorter context slice for inline auto-verification (keeps latency manageable)
    verify_ref_context = context[:20_000]

    async def stream() -> AsyncGenerator[str, None]:
        yield _sse({"type": "start", "total": len(output_types)})

        proj.status = "generating"
        db.commit()

        for idx, output_type in enumerate(output_types):
            yield _sse({"type": "output_start", "output_type": output_type, "index": idx,
                        "label": _OUTPUT_LABELS.get(output_type, output_type)})
            try:
                user_prompt = _build_type_prompt(output_type, proj)
                full_content = ""
                async for chunk in provider.stream_chat(
                    [{"role": "user", "content": user_prompt}],
                    system_prompt=system_prompt,
                ):
                    full_content += chunk
                    yield _sse({"type": "chunk", "chunk": chunk, "output_type": output_type})

                citations = _extract_citations_from_content(full_content, refs)
                quality   = _compute_quality(full_content, citations, refs)

                # Upsert output
                existing = db.query(EducationOutput).filter_by(
                    project_id=project_id, output_type=output_type
                ).first()
                if existing and body.regenerate:
                    existing.content           = full_content
                    existing.citations         = citations
                    existing.quality_issues    = quality["quality_issues"]
                    existing.quality_score     = quality["quality_score"]
                    existing.technical_accuracy_score = quality["technical_accuracy_score"]
                    existing.source_coverage_score    = quality["source_coverage_score"]
                    existing.citation_score           = quality["citation_score"]
                    existing.lo_alignment_score       = quality["lo_alignment_score"]
                    existing.title                    = _OUTPUT_LABELS.get(output_type, output_type)
                    existing.approved                 = False
                    db.commit()
                    saved_out = existing
                    out_id = existing.id
                elif not existing:
                    new_out = EducationOutput(
                        project_id=project_id,
                        output_type=output_type,
                        title=_OUTPUT_LABELS.get(output_type, output_type),
                        content=full_content,
                        citations=citations,
                        quality_issues=quality["quality_issues"],
                        quality_score=quality["quality_score"],
                        technical_accuracy_score=quality["technical_accuracy_score"],
                        source_coverage_score=quality["source_coverage_score"],
                        citation_score=quality["citation_score"],
                        lo_alignment_score=quality["lo_alignment_score"],
                    )
                    db.add(new_out)
                    db.commit()
                    db.refresh(new_out)
                    saved_out = new_out
                    out_id = new_out.id
                else:
                    saved_out = existing
                    out_id = existing.id

                yield _sse({
                    "type": "output_done",
                    "output_type": output_type,
                    "output_id": out_id,
                    "quality_score": quality["quality_score"],
                    "citation_count": len(citations),
                })

                # ── Inline auto-verification ──────────────────────────────
                # Runs immediately after each output is saved.  This wires
                # reference-drift detection into the generation pipeline so
                # flags are available before the user ever opens the Review
                # step, not only when they click "Verify" manually.
                yield _sse({
                    "type": "output_verifying",
                    "output_type": output_type,
                    "output_id": out_id,
                })
                try:
                    vr = await _run_verification_call(
                        full_content, verify_ref_context, provider, max_paragraphs=20
                    )
                    _apply_verification_to_output(saved_out, vr)
                    db.commit()
                    yield _sse({
                        "type":          "output_verified",
                        "output_type":   output_type,
                        "output_id":     out_id,
                        "flag_count":    len(vr.get("flagged", [])),
                        "flagged":       vr.get("flagged", []),
                        "summary":       vr.get("summary", ""),
                        "quality_score": saved_out.quality_score,
                    })
                except Exception as ve:
                    # Verification failure is non-blocking — log it and move on.
                    log.warning(
                        "Auto-verification non-blocking failure for %s: %s",
                        output_type, ve,
                    )
                    yield _sse({
                        "type":        "output_verified",
                        "output_type": output_type,
                        "output_id":   out_id,
                        "flag_count":  -1,
                        "error":       "Verification could not run for this output.",
                    })

            except Exception as exc:
                log.error("Generation error for %s: %s", output_type, exc)
                yield _sse({"type": "output_error", "output_type": output_type, "error": str(exc)})

        proj.status = "complete"
        # Compute overall quality
        all_outputs = db.query(EducationOutput).filter_by(project_id=project_id).all()
        if all_outputs:
            scores = [o.quality_score for o in all_outputs if o.quality_score is not None]
            proj.quality_score = int(sum(scores) / len(scores)) if scores else None
        db.commit()
        yield _sse({"type": "done", "project_id": project_id})

    return StreamingResponse(
        stream(),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


# ── Output routes ─────────────────────────────────────────────────────────────

@router.get("/edu/projects/{project_id}/outputs")
def list_outputs(project_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    _get_project(db, project_id, user["id"])
    outputs = db.query(EducationOutput).filter_by(project_id=project_id).all()
    return {"outputs": [_output_dict(o) for o in outputs]}


@router.get("/edu/projects/{project_id}/outputs/{out_id}")
def get_output(project_id: str, out_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    _get_project(db, project_id, user["id"])
    out = _get_output(db, out_id, project_id)
    return _output_dict(out, include_content=True)


@router.patch("/edu/projects/{project_id}/outputs/{out_id}")
def patch_output(
    project_id: str, out_id: str, body: OutputPatch,
    request: Request, db: Session = Depends(get_db)
):
    user = require_auth(request)
    _get_project(db, project_id, user["id"])
    out = _get_output(db, out_id, project_id)
    if body.content is not None:
        out.content = body.content
    if body.title is not None:
        out.title = body.title
    if body.approved is not None:
        out.approved = body.approved
    db.commit()
    return _output_dict(out)


# ── Verification helper ───────────────────────────────────────────────────────

async def _run_verification_call(
    content: str,
    ref_context: str,
    provider: Any,
    max_paragraphs: int = 30,
) -> dict:
    """
    Run a citation/grounding verification AI call.

    Segments content into paragraphs (filtering blanks, the same rule used
    everywhere), sends them with the reference context, and returns a parsed
    result dict: {flagged, verified_citation_score,
    verified_technical_accuracy_score, summary}.

    Raises ValueError (not HTTPException) on parse failure so callers decide
    whether to propagate as HTTP 502 or swallow as a warning.
    """
    # ── Segment — must match the frontend filter exactly ──
    paragraphs = [p.strip() for p in content.split("\n\n") if p.strip()]
    para_sample = paragraphs[:max_paragraphs]

    numbered_paras = "\n\n".join(
        f"[P{i}] {p[:600]}" for i, p in enumerate(para_sample)
    )

    verify_system = (
        "You are a precise technical training quality auditor. "
        "Return only valid JSON — no markdown, no prose outside the JSON object."
    )

    verify_prompt = (
        "You are auditing generated training content against its source reference material.\n\n"
        "REFERENCE MATERIAL (authoritative source):\n"
        + ref_context
        + "\n\n"
        "GENERATED CONTENT (each paragraph labelled [P0], [P1], …):\n"
        + numbered_paras
        + "\n\n"
        "For every paragraph, check:\n"
        "1. Does it make specific technical claims (values, part numbers, procedures,\n"
        "   specifications) WITHOUT a [Source: filename, p.N] citation?\n"
        "2. Does any specific technical value CONTRADICT what the reference states?\n"
        "3. Does it introduce technical details NOT present anywhere in the reference?\n\n"
        "ONLY flag paragraphs with genuine issues — skip introductory, transitional, or\n"
        "purely structural text that makes no technical claims.\n\n"
        "Return a JSON object with EXACTLY this structure:\n"
        "{\n"
        '  "flagged": [\n'
        "    {\n"
        '      "paragraph_index": <integer>,\n'
        '      "concern": "<concise description of the specific issue>",\n'
        '      "severity": "error" | "warning",\n'
        '      "paragraph_preview": "<first 120 chars of the paragraph>"\n'
        "    }\n"
        "  ],\n"
        '  "verified_citation_score": <integer 0-100>,\n'
        '  "verified_technical_accuracy_score": <integer 0-100>,\n'
        '  "summary": "<one sentence overall verdict>"\n'
        "}"
    )

    full_response = ""
    async for chunk in provider.stream_chat(
        [{"role": "user", "content": verify_prompt}],
        system_prompt=verify_system,
    ):
        full_response += chunk

    # Parse JSON — fail CLOSED: raise ValueError on any parse problem.
    json_match = re.search(r'\{.*\}', full_response, re.DOTALL)
    if not json_match:
        raise ValueError(
            f"No JSON in verification response. Raw (first 300 chars): {full_response[:300]!r}"
        )
    try:
        result: dict = json.loads(json_match.group())
    except json.JSONDecodeError as exc:
        raise ValueError(f"Malformed JSON in verification response: {exc}") from exc

    if "flagged" not in result:
        raise ValueError(
            f"'flagged' key missing from verification response. Keys: {list(result.keys())}"
        )

    return result


def _apply_verification_to_output(
    out: Any,
    result: dict,
) -> None:
    """
    Persist verification results onto an EducationOutput ORM object in-place.
    Does NOT commit — caller must commit.
    """
    flagged: list[dict] = result.get("flagged", [])
    v_cite  = result.get("verified_citation_score")
    v_tech  = result.get("verified_technical_accuracy_score")

    # Update verified subscores
    if isinstance(v_cite, (int, float)):
        out.citation_score = int(v_cite)
    if isinstance(v_tech, (int, float)):
        out.technical_accuracy_score = int(v_tech)

    # Recompute overall quality with verified subscores
    cite = out.citation_score or 0
    tech = out.technical_accuracy_score or 0
    cov  = out.source_coverage_score or 0
    lo   = out.lo_alignment_score or 80
    out.quality_score = int(cite * 0.35 + cov * 0.25 + tech * 0.25 + lo * 0.15)

    # Merge flags into quality_issues, replacing any previous verify_flag entries
    existing_issues = [
        i for i in (out.quality_issues or [])
        if i.get("type") != "verify_flag"
    ]
    for flag in flagged:
        existing_issues.append({
            "type":              "verify_flag",
            "severity":          flag.get("severity", "warning"),
            "message":           flag.get("concern", ""),
            "paragraph_index":   flag.get("paragraph_index"),
            "paragraph_preview": flag.get("paragraph_preview", ""),
        })
    out.quality_issues = existing_issues


@router.post("/edu/projects/{project_id}/outputs/{out_id}/verify")
async def verify_output(
    project_id: str, out_id: str,
    request: Request, db: Session = Depends(get_db),
):
    """
    Manual second-pass citation verifier.

    Splits the output into paragraphs (filtering blanks, matching the frontend
    indexing rule) and asks the AI to cross-check each one against the
    extracted reference text.  Flagged paragraphs are persisted in
    quality_issues and returned to the frontend.  citation_score and
    technical_accuracy_score are updated with AI-verified values.
    """
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    out  = _get_output(db, out_id, project_id)

    provider = provider_registry.get_active()
    if not provider:
        raise HTTPException(503, "No AI provider available")

    refs = [r for r in proj.references if r.status == "done"]
    if not refs:
        raise HTTPException(400, "No reference documents available for verification")

    ref_context = _build_reference_context(refs)[:25_000]
    paragraphs  = [p.strip() for p in out.content.split("\n\n") if p.strip()]

    try:
        result = await _run_verification_call(
            out.content, ref_context, provider, max_paragraphs=50
        )
    except ValueError as exc:
        log.error("Manual verify parse failure for output %s: %s", out_id, exc)
        raise HTTPException(
            502,
            detail={
                "error": "verification_parse_failed",
                "message": (
                    "The AI did not return a parseable verification response. "
                    "Please try again — if the problem persists, check the AI provider."
                ),
            },
        ) from exc

    _apply_verification_to_output(out, result)
    db.commit()

    flagged = result.get("flagged", [])
    return {
        "output_id":                         out_id,
        "flagged":                           flagged,
        "verified_citation_score":           out.citation_score,
        "verified_technical_accuracy_score": out.technical_accuracy_score,
        "quality_score":                     out.quality_score,
        "summary":                           result.get("summary", ""),
        "paragraph_count":                   len(paragraphs[:50]),
    }


@router.post("/edu/projects/{project_id}/outputs/{out_id}/regenerate")
async def regenerate_output(
    project_id: str, out_id: str,
    request: Request, db: Session = Depends(get_db),
):
    user = require_auth(request)
    proj = _get_project(db, project_id, user["id"])
    out  = _get_output(db, out_id, project_id)

    provider = provider_registry.get_active()
    if not provider:
        raise HTTPException(503, "No AI provider available")

    refs    = [r for r in proj.references if r.status == "done"]
    context = _build_reference_context(refs)
    sys_p   = _build_system_prompt(proj, context)
    usr_p   = _build_type_prompt(out.output_type, proj)
    verify_ref_context = context[:20_000]

    async def stream() -> AsyncGenerator[str, None]:
        full = ""
        async for chunk in provider.stream_chat(
            [{"role": "user", "content": usr_p}], system_prompt=sys_p
        ):
            full += chunk
            yield _sse({"type": "chunk", "chunk": chunk})

        citations = _extract_citations_from_content(full, refs)
        quality   = _compute_quality(full, citations, refs)
        out.content           = full
        out.citations         = citations
        out.quality_issues    = quality["quality_issues"]
        out.quality_score     = quality["quality_score"]
        out.technical_accuracy_score = quality["technical_accuracy_score"]
        out.source_coverage_score    = quality["source_coverage_score"]
        out.citation_score           = quality["citation_score"]
        out.lo_alignment_score       = quality["lo_alignment_score"]
        out.approved                 = False
        db.commit()
        yield _sse({"type": "done", "output_id": out.id, "quality_score": quality["quality_score"]})

        # Auto-verify after regeneration — same as in generate_outputs
        yield _sse({"type": "output_verifying", "output_id": out.id})
        try:
            vr = await _run_verification_call(
                full, verify_ref_context, provider, max_paragraphs=20
            )
            _apply_verification_to_output(out, vr)
            db.commit()
            yield _sse({
                "type":          "output_verified",
                "output_id":     out.id,
                "flag_count":    len(vr.get("flagged", [])),
                "flagged":       vr.get("flagged", []),
                "summary":       vr.get("summary", ""),
                "quality_score": out.quality_score,
            })
        except Exception as ve:
            log.warning("Auto-verify failed on regen for %s: %s", out.id, ve)
            yield _sse({
                "type":       "output_verified",
                "output_id":  out.id,
                "flag_count": -1,
                "error":      "Verification could not run for this output.",
            })

    return StreamingResponse(
        stream(), media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )


# ── Quality report ────────────────────────────────────────────────────────────

@router.get("/edu/projects/{project_id}/quality-report")
def quality_report(project_id: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj    = _get_project(db, project_id, user["id"])
    outputs = db.query(EducationOutput).filter_by(project_id=project_id).all()

    output_reports = []
    for o in outputs:
        output_reports.append({
            "output_type": o.output_type,
            "title":       o.title,
            "quality_score": o.quality_score,
            "technical_accuracy_score": o.technical_accuracy_score,
            "source_coverage_score":    o.source_coverage_score,
            "citation_score":           o.citation_score,
            "lo_alignment_score":       o.lo_alignment_score,
            "quality_issues": o.quality_issues or [],
            "citation_count": len(o.citations or []),
            "approved": o.approved,
        })

    scores = [o.quality_score for o in outputs if o.quality_score is not None]
    overall = int(sum(scores) / len(scores)) if scores else None

    all_issues = [issue for o in outputs for issue in (o.quality_issues or [])]
    errors   = [i for i in all_issues if i.get("severity") == "error"]
    warnings = [i for i in all_issues if i.get("severity") == "warning"]

    return {
        "project_id":    project_id,
        "overall_score": overall,
        "output_count":  len(outputs),
        "approved_count":sum(1 for o in outputs if o.approved),
        "errors":        len(errors),
        "warnings":      len(warnings),
        "all_issues":    all_issues,
        "outputs":       output_reports,
    }


# ── Export ────────────────────────────────────────────────────────────────────

@router.get("/edu/projects/{project_id}/export/{fmt}")
def export_project(project_id: str, fmt: str, request: Request, db: Session = Depends(get_db)):
    user = require_auth(request)
    proj    = _get_project(db, project_id, user["id"])
    outputs = db.query(EducationOutput).filter_by(project_id=project_id).all()
    refs    = proj.references

    fmt = fmt.lower()

    if fmt == "zip":
        return _export_zip(proj, outputs, refs)
    elif fmt == "docx":
        return _export_docx(proj, outputs)
    elif fmt == "pptx":
        pptx_out = next((o for o in outputs if o.output_type == "pptx_presentation"), None)
        content = pptx_out.content if pptx_out else "\n\n".join(o.content for o in outputs[:1])
        return _export_pptx(proj, content, refs=list(refs))
    elif fmt == "txt":
        return _export_txt(proj, outputs)
    elif fmt == "csv":
        return _export_csv(proj, outputs)
    else:
        raise HTTPException(400, f"Unsupported format: {fmt}")


def _export_txt(proj: EducationProject, outputs: list[EducationOutput]) -> Response:
    lines = [f"# {proj.title}", f"Course: {proj.course_title}", "=" * 60, ""]
    for o in outputs:
        lines += [f"\n{'='*60}", f"## {o.title}", f"{'='*60}\n", o.content, ""]
    content = "\n".join(lines)
    return Response(
        content=content.encode("utf-8"),
        media_type="text/plain",
        headers={"Content-Disposition": f'attachment; filename="{proj.title}.txt"'},
    )


def _export_csv(proj: EducationProject, outputs: list[EducationOutput]) -> Response:
    import csv, io as _io
    buf = _io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Question", "A", "B", "C", "D", "Correct", "Explanation", "Source", "Difficulty"])

    for o in outputs:
        if o.output_type not in ("quiz", "formal_exam", "knowledge_check"):
            continue
        q_re = re.compile(
            r'\*\*Q\d+\.\*\*\s+(.+?)[\n\r]+'
            r'A\)\s+(.+?)[\n\r]+B\)\s+(.+?)[\n\r]+C\)\s+(.+?)[\n\r]+D\)\s+(.+?)[\n\r]+'
            r'\*\*(?:Correct\s+)?Answer:\*\*\s+([A-D])[^\n]*(?:\n\*\*Explanation:\*\*\s+(.+?))?(?=\n\n|\Z)',
            re.DOTALL | re.I,
        )
        for m in q_re.finditer(o.content):
            w.writerow([
                m.group(1).strip(), m.group(2).strip(), m.group(3).strip(),
                m.group(4).strip(), m.group(5).strip(), m.group(6).strip(),
                m.group(7).strip() if m.group(7) else "", "", ""
            ])

    return Response(
        content=buf.getvalue().encode("utf-8"),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{proj.title}_questions.csv"'},
    )


def _export_docx(proj: EducationProject, outputs: list[EducationOutput]) -> Response:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title page
    title_para = doc.add_heading(proj.title, 0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if proj.course_title:
        doc.add_paragraph(f"Course: {proj.course_title}")
    if proj.equipment_manufacturer:
        doc.add_paragraph(f"Equipment: {proj.equipment_manufacturer} {proj.equipment_model}")
    doc.add_paragraph(f"Audience: {proj.audience}  |  Level: {proj.level}")
    doc.add_paragraph(f"Generated: {datetime.now(timezone.utc).strftime('%Y-%m-%d')}")
    doc.add_page_break()

    for o in outputs:
        doc.add_heading(o.title or o.output_type.replace("_", " ").title(), level=1)
        # Convert markdown to docx paragraphs
        for line in o.content.splitlines():
            stripped = line.strip()
            if not stripped:
                doc.add_paragraph("")
                continue
            if stripped.startswith("# "):
                doc.add_heading(stripped[2:], level=1)
            elif stripped.startswith("## "):
                doc.add_heading(stripped[3:], level=2)
            elif stripped.startswith("### "):
                doc.add_heading(stripped[4:], level=3)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                p = doc.add_paragraph(stripped[2:], style="List Bullet")
            elif re.match(r'^\d+\.', stripped):
                doc.add_paragraph(stripped, style="List Number")
            else:
                doc.add_paragraph(stripped)
        doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return Response(
        content=buf.read(),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{proj.title}.docx"'},
    )


def _parse_slide_content(content: str) -> list[dict]:
    """
    Parse GPT markdown slide output into structured slide dicts for pptx_gen.build_pptx.

    Expected format per slide:
        ## Slide N: Title
        **Type:** content
        **Bullets:**
        - bullet text
        **Image:** FIGURE_N | none
        **Speaker Notes:** ...
        **Source:** [Source: FILE, p.N]
    """
    slides: list[dict] = []
    current: dict | None = None
    in_bullets = False

    for line in content.splitlines():
        stripped = line.strip()

        # New slide marker — ## Slide N: Title  OR  ## Title (fallback)
        m_slide = re.match(r'^## Slide\s+\d+[:.]\s*(.+)', stripped, re.I)
        m_section = re.match(r'^## (.+)', stripped)
        if m_slide or (m_section and stripped.lower() not in
                       ("## bullets:", "## speaker notes:", "## source:")):
            if current:
                slides.append(current)
            title = (m_slide or m_section).group(1).strip()  # type: ignore[union-attr]
            # Strip trailing ** noise from GPT markdown
            title = re.sub(r'\*+', '', title).strip()
            current = {
                "type": "content",
                "title": title[:100],
                "bullets": [],
                "speaker_notes": "",
                "source_pages": [],
                "image_tag": None,
                "is_visible": True,
            }
            in_bullets = False
            continue

        if current is None:
            continue

        if stripped.startswith("**Type:**"):
            raw_type = stripped[9:].strip().lower()
            type_map = {
                "image_content":   "image_content",
                "full_image":      "full_image",
                "warning":         "content",
                "procedure":       "procedure",
                "practical":       "practical",
                "quiz":            "quiz",
                "knowledge_check": "quiz",
                "objectives":      "objectives",
                "section":         "section",
                "tracker":         "tracker",
                "lab":             "lab",
                "title":           "title",
                "agenda":          "agenda",
                "summary":         "summary",
                "references":      "references",
            }
            current["type"] = type_map.get(raw_type, "content")
            in_bullets = False

        elif stripped.lower().startswith("**bullets:**") or stripped.lower() == "**content:**":
            in_bullets = True

        elif stripped.startswith("**Image:**"):
            img_val = stripped[10:].strip()
            if img_val.lower() not in ("none", "", "n/a"):
                current["image_tag"] = img_val.upper()
                if current["type"] == "content":
                    current["type"] = "image_content"
            in_bullets = False

        elif stripped.startswith("**Speaker Notes:**"):
            current["speaker_notes"] = stripped[18:].strip()
            in_bullets = False

        elif stripped.startswith("**Source:**"):
            src = stripped[11:].strip()
            pages = re.findall(r'p\.?(\d+)', src)
            current["source_pages"] = [int(p) for p in pages[:8]]
            in_bullets = False

        elif stripped.startswith("**Answer:**") or re.match(r'\*\*Correct\s*Answer', stripped, re.I):
            m_ans = re.search(r'\b([A-D])\b', stripped)
            if m_ans:
                current["answer_index"] = ord(m_ans.group(1).upper()) - ord('A')
            in_bullets = False

        elif in_bullets and (stripped.startswith("- ") or stripped.startswith("* ")):
            current["bullets"].append(stripped[2:])

        elif in_bullets and (line.startswith("  ") or line.startswith("\t")):
            current["bullets"].append(stripped)

    if current:
        slides.append(current)

    return slides


def _export_zip(
    proj: EducationProject,
    outputs: list[EducationOutput],
    refs: list[EducationReference],
) -> Response:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Index
        index_lines = [
            f"# {proj.title}",
            f"Generated: {datetime.now(timezone.utc).isoformat()}",
            f"Course: {proj.course_title}",
            f"Equipment: {proj.equipment_manufacturer} {proj.equipment_model}",
            f"Audience: {proj.audience} | Level: {proj.level}",
            "",
            "## Contents",
        ]
        for o in outputs:
            index_lines.append(f"- {o.title}: {o.output_type}.txt")
        zf.writestr("00_index.txt", "\n".join(index_lines))

        # Each output as txt
        for o in outputs:
            safe_name = re.sub(r'[^a-z0-9_]', '_', o.output_type)
            zf.writestr(f"outputs/{safe_name}.txt", o.content)

        # Quality report
        qr_lines = [f"# Quality Report — {proj.title}", ""]
        for o in outputs:
            qr_lines += [
                f"## {o.title}",
                f"- Overall: {o.quality_score}/100",
                f"- Citation score: {o.citation_score}/100",
                f"- Source coverage: {o.source_coverage_score}/100",
                f"- Technical accuracy: {o.technical_accuracy_score}/100",
                f"- LO alignment: {o.lo_alignment_score}/100",
                "",
            ]
            for issue in (o.quality_issues or []):
                qr_lines.append(f"  [{issue.get('severity','warn').upper()}] {issue.get('message','')}")
            qr_lines.append("")
        zf.writestr("quality_report.txt", "\n".join(qr_lines))

        # References summary
        ref_lines = [f"# Reference Files — {proj.title}", ""]
        for r in refs:
            ref_lines += [
                f"## {r.filename}",
                f"- Type: {r.file_type} | Pages: {r.page_count} | Words: {r.word_count}",
                f"- Role: {r.role}",
                f"- Sections: {r.section_count} | Procedures: {r.procedure_count} | Warnings: {r.warning_count}",
                "",
            ]
        zf.writestr("references.txt", "\n".join(ref_lines))

    buf.seek(0)
    safe_title = re.sub(r'[^a-z0-9_]', '_', proj.title.lower())[:40]
    return Response(
        content=buf.read(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}_package.zip"'},
    )


def _extract_reference_figures(refs: list) -> dict[str, bytes]:
    """
    Extract embedded images from PDF references using PyMuPDF.
    Returns {FIGURE_N: png_bytes} in the same order as the figure catalogue
    presented to GPT in _build_reference_context.
    Only images larger than 8 KB are included (filters out tiny icons/logos).
    """
    try:
        import fitz  # PyMuPDF
    except ImportError:
        log.warning("PyMuPDF not available — PPTX figures disabled")
        return {}

    figures: dict[str, bytes] = {}
    fig_n = 1

    for ref in refs:
        if not ref.raw_bytes or ref.file_type != "pdf":
            continue
        try:
            doc = fitz.open(stream=ref.raw_bytes, filetype="pdf")
            for page in doc:
                for img_info in page.get_images(full=True):
                    xref = img_info[0]
                    try:
                        base_img = doc.extract_image(xref)
                        img_bytes = base_img["image"]
                        if len(img_bytes) < 8_000:
                            continue  # skip tiny icons

                        # Convert to PNG for python-pptx compatibility
                        if base_img["ext"].lower() not in ("png", "jpeg", "jpg"):
                            try:
                                from PIL import Image as PILImage
                                pil = PILImage.open(io.BytesIO(img_bytes))
                                buf = io.BytesIO()
                                pil.save(buf, "PNG")
                                img_bytes = buf.getvalue()
                            except Exception:
                                continue  # skip unconvertible formats

                        figures[f"FIGURE_{fig_n}"] = img_bytes
                        fig_n += 1
                    except Exception:
                        continue
            doc.close()
        except Exception as exc:
            log.warning("Figure extraction failed for %s: %s", ref.filename, exc)

    log.info("Extracted %d figures from references for PPTX", len(figures))
    return figures


def _export_pptx(proj: EducationProject, content: str,
                 refs: list | None = None) -> Response:
    """
    Build a full Rapiscan-style PPTX using pptx_gen.build_pptx.
    Parses the GPT markdown slide content, extracts figures from reference PDFs,
    then delegates rendering to the professional slide builder.
    """
    from api.utils.pptx_gen import build_pptx

    slides = _parse_slide_content(content)

    # Fall back to a simple title slide if parsing yields nothing
    if not slides:
        slides = [{
            "type": "title",
            "title": proj.course_title or proj.title,
            "bullets": [
                f"{proj.equipment_manufacturer} {proj.equipment_model}",
                f"Audience: {proj.audience}  |  Level: {proj.level}",
            ],
            "speaker_notes": "",
            "source_pages": [],
            "is_visible": True,
        }]

    # Extract figures from reference PDFs
    figures: dict[str, bytes] = {}
    if refs and proj.include_images:
        figures = _extract_reference_figures(refs)

    project_meta = {
        "course_title":    proj.course_title or proj.title,
        "manufacturer":    proj.equipment_manufacturer,
        "equipment_model": proj.equipment_model,
        "audience":        proj.audience,
        "training_type":   proj.level,
        "language":        proj.language,
        "manual_filename": refs[0].filename if refs else "",
    }

    pptx_bytes = build_pptx(project_meta, slides, version="combined", images=figures)

    safe_name = re.sub(r'[^a-z0-9_\-]', '_', (proj.title or "course").lower())[:50]
    return Response(
        content=pptx_bytes,
        media_type="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        headers={"Content-Disposition": f'attachment; filename="{safe_name}.pptx"'},
    )


def _export_zip(
    proj: EducationProject,
    outputs: list[EducationOutput],
    refs: list[EducationReference],
) -> Response:
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        # Index
        index_lines = [
            f"# {proj.title}",
            f"Generated: {datetime.now(timezone.utc).isoformat()}",
            f"Course: {proj.course_title}",
            f"Equipment: {proj.equipment_manufacturer} {proj.equipment_model}",
            f"Audience: {proj.audience} | Level: {proj.level}",
            "",
            "## Contents",
        ]
        for o in outputs:
            index_lines.append(f"- {o.title}: {o.output_type}.txt")
        zf.writestr("00_index.txt", "\n".join(index_lines))

        # Each output as txt
        for o in outputs:
            safe_name = re.sub(r'[^a-z0-9_]', '_', o.output_type)
            zf.writestr(f"outputs/{safe_name}.txt", o.content)

        # Quality report
        qr_lines = [f"# Quality Report — {proj.title}", ""]
        for o in outputs:
            qr_lines += [
                f"## {o.title}",
                f"- Overall: {o.quality_score}/100",
                f"- Citation score: {o.citation_score}/100",
                f"- Source coverage: {o.source_coverage_score}/100",
                f"- Technical accuracy: {o.technical_accuracy_score}/100",
                f"- LO alignment: {o.lo_alignment_score}/100",
                "",
            ]
            for issue in (o.quality_issues or []):
                qr_lines.append(f"  [{issue.get('severity','warn').upper()}] {issue.get('message','')}")
            qr_lines.append("")
        zf.writestr("quality_report.txt", "\n".join(qr_lines))

        # References summary
        ref_lines = [f"# Reference Files — {proj.title}", ""]
        for r in refs:
            ref_lines += [
                f"## {r.filename}",
                f"- Type: {r.file_type} | Pages: {r.page_count} | Words: {r.word_count}",
                f"- Role: {r.role}",
                f"- Sections: {r.section_count} | Procedures: {r.procedure_count} | Warnings: {r.warning_count}",
                "",
            ]
        zf.writestr("references.txt", "\n".join(ref_lines))

    buf.seek(0)
    safe_title = re.sub(r'[^a-z0-9_]', '_', proj.title.lower())[:40]
    return Response(
        content=buf.read(),
        media_type="application/zip",
        headers={"Content-Disposition": f'attachment; filename="{safe_title}_package.zip"'},
    )


# ── Slide approval & correction endpoints ──────────────────────────────────────

class SlideApproveBody(BaseModel):
    approved: bool = True


class SlideCorrectionBody(BaseModel):
    original_content: dict = {}
    corrected_content: dict
    correction_reason: str = ""
    corrected_by: str = ""


@router.post("/edu/{project_id}/slides/{output_id}/approve")
def approve_education_output(
    project_id: str,
    output_id: str,
    body: SlideApproveBody,
    request: Request,
    db: Session = Depends(get_db),
):
    """Approve or un-approve an Education Studio output (authenticated users only)."""
    from datetime import datetime, timezone

    user = require_auth(request)  # raises 401 if not logged in

    proj = db.query(EducationProject).filter(EducationProject.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    # Ownership check: project owner or admin
    user_id = str(user.get("id") or "")
    if hasattr(proj, "user_id") and proj.user_id and proj.user_id != user_id:
        raise HTTPException(status_code=403, detail="Not authorised to modify this project")

    out = db.query(EducationOutput).filter(
        EducationOutput.id == output_id,
        EducationOutput.project_id == project_id,
    ).first()
    if not out:
        raise HTTPException(status_code=404, detail="Output not found")

    out.approved = body.approved
    out.approved_at = datetime.now(timezone.utc) if body.approved else None
    db.commit()
    return {"ok": True, "output_id": output_id, "approved": out.approved}


@router.post("/edu/{project_id}/slides/{output_id}/correct")
def correct_education_output(
    project_id: str,
    output_id: str,
    body: SlideCorrectionBody,
    request: Request,
    db: Session = Depends(get_db),
):
    """Record an instructor correction for an Education Studio output (authenticated users only)."""
    from api.db.models import SlideCorrection

    user = require_auth(request)  # raises 401 if not logged in

    proj = db.query(EducationProject).filter(EducationProject.id == project_id).first()
    if not proj:
        raise HTTPException(status_code=404, detail="Project not found")

    # Ownership check
    user_id = str(user.get("id") or "")
    if hasattr(proj, "user_id") and proj.user_id and proj.user_id != user_id:
        raise HTTPException(status_code=403, detail="Not authorised to modify this project")

    out = db.query(EducationOutput).filter(
        EducationOutput.id == output_id,
        EducationOutput.project_id == project_id,
    ).first()
    if not out:
        raise HTTPException(status_code=404, detail="Output not found")

    original = body.original_content or {"content": out.content}

    correction = SlideCorrection(
        slide_id=output_id,
        slide_type="education",
        original_content=original,
        corrected_content=body.corrected_content,
        correction_reason=body.correction_reason,
        corrected_by=body.corrected_by or user.get("name") or user.get("email") or "instructor",
    )
    db.add(correction)

    # Apply correction to the output content if corrected text was provided
    new_content = body.corrected_content.get("content")
    if new_content and isinstance(new_content, str):
        out.content = new_content

    db.commit()
    return {"ok": True, "correction_id": correction.id}
