"""
Training Course Report — DOCX template engine.

Populates the canonical "Training Course Report v3 - Blank Copy.docx"
template (stored permanently under backend/templates/training_reports/)
in place with python-docx, so the original page size, margins, header/
footer/logo, fonts, table borders and pagination are preserved exactly —
nothing is rebuilt from scratch.

Narrative sections are replaced as whole paragraph blocks (heading found
by exact text match, body block = every paragraph up to the next known
heading) rather than token-patched, so no bracketed placeholder such as
"(Days)" or "(Select System)" can survive into the final document.
"""
from __future__ import annotations

import copy
import io
import os
import re
from typing import Optional

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))), "templates", "training_reports")
TEMPLATE_PATH = os.path.join(TEMPLATE_DIR, "training_course_report_v3_blank.docx")

# Every literal placeholder token that appears in the blank template.
# Used only for the post-generation "no placeholder left behind" check.
KNOWN_PLACEHOLDER_TOKENS = [
    "(Days)", "(Type of Training)", "(System Model)", "(Customer Site/ Location)",
    "(Customer Site/Location)", "(Name)", "(English, Latin etc)", "(English)",
    "(Location)", "(….)", "(…)", "(Expand as required)", "(Select System)",
    "(Example -",
]

# Section headings in document order — used both to locate a section's
# heading paragraph and as the "stop" set when collecting its body block.
_HEADINGS = [
    "Reference Documents",
    "Background",
    "Instructor",
    "Training Delivery Language",
    "Training Location",
    "Students",
    "Schedule of Training",
    "Course Completion and Certification",
    "Additional Notes and QA Issues",
    "After Training follow up Plan:",
]

# section_key -> exact heading text in the template
SECTION_HEADINGS = {
    "reference_documents": "Reference Documents",
    "background": "Background",
    "instructor": "Instructor",
    "delivery_language": "Training Delivery Language",
    "training_location": "Training Location",
    "students": "Students",
    "schedule": "Schedule of Training",
    "completion_certification": "Course Completion and Certification",
    "additional_notes": "Additional Notes and QA Issues",
    "follow_up_plan": "After Training follow up Plan:",
}


def load_template():
    from docx import Document
    if not os.path.exists(TEMPLATE_PATH):
        raise FileNotFoundError(f"Training report template not found at {TEMPLATE_PATH}")
    return Document(TEMPLATE_PATH)


def fmt_num(value) -> str:
    """Render 5.0 as "5" but keep 5.5 as "5.5" — value entered via a number
    input always arrives as a float even when the user meant a whole number."""
    if isinstance(value, float) and value.is_integer():
        return str(int(value))
    return str(value)


def _norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "")).strip().rstrip("*").strip()


def _heading_paragraph(doc, heading_text: str):
    target = _norm(heading_text)
    for p in doc.paragraphs:
        if _norm(p.text) == target:
            return p
    return None


def _section_block_paragraphs(doc, heading_text: str) -> list:
    """All body paragraphs belonging to one section: everything between its
    heading and the next known heading (or end of body)."""
    paras = doc.paragraphs
    stop_texts = {_norm(h) for h in _HEADINGS} | {"TRAINING REPORT"}
    start_idx = None
    for i, p in enumerate(paras):
        if _norm(p.text) == _norm(heading_text):
            start_idx = i
            break
    if start_idx is None:
        return []
    block = []
    for p in paras[start_idx + 1:]:
        if _norm(p.text) in stop_texts:
            break
        block.append(p)
    return block


def _clear_paragraph_runs(paragraph) -> None:
    for run in list(paragraph.runs):
        run.text = ""
    # drop every run but the first so no stray empty runs remain
    for run_elm in list(paragraph._p.findall("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"))[1:]:
        paragraph._p.remove(run_elm)


def _set_rtl_paragraph(paragraph) -> None:
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in paragraph.runs:
        rPr = run._element.get_or_add_rPr()
        rPr.append(OxmlElement("w:rtl"))


def replace_section(doc, section_key: str, text: Optional[str], is_rtl: bool = False) -> None:
    """Replace a section's whole body block with `text` (or blank it if
    `text` is falsy — heading stays, body is simply removed)."""
    heading_text = SECTION_HEADINGS[section_key]
    heading_para = _heading_paragraph(doc, heading_text)
    if heading_para is None:
        return
    block = _section_block_paragraphs(doc, heading_text)
    if not block:
        return
    first = block[0]
    _clear_paragraph_runs(first)
    if text and text.strip():
        run = first.add_run(text.strip())
        if is_rtl:
            _set_rtl_paragraph(first)
    for extra in block[1:]:
        extra._p.getparent().remove(extra._p)


def _label_key(text: str) -> str:
    return _norm(text).lower()


_INFO_LABEL_MAP = {
    "training level": "training_level",
    "dates of training": "dates_of_training",
    "instructor names": "instructor",
    "instructor name(s)": "instructor",
    "project code": "project_code",
    "system type": "system_type",
    "system serial no.": "system_serial_no",
    "site of operation": "site_of_operation",
    "customer": "customer",
    "contract reference": "contract_reference",
}


def fill_info_table(doc, values: dict) -> None:
    table = doc.tables[0]
    for row in table.rows:
        label = _label_key(row.cells[0].text)
        if label == "course attendees":
            row.cells[1].text = f"Total {values.get('attendees_total', 0)}"
            row.cells[2].text = (
                f"Pass {values.get('attendees_pass', 0)}   "
                f"Failed {values.get('attendees_fail', 0)}   "
                f"Certified {values.get('attendees_certified', 0)}"
            )
            continue
        key = _INFO_LABEL_MAP.get(label)
        if key and key in values:
            row.cells[1].text = str(values.get(key) or "")


def fill_doc_ref_table(doc, refs: list[dict]) -> None:
    table = doc.tables[1]
    rows = [r for r in refs if (r.get("reference_number") or "").strip() or (r.get("title") or "").strip()]
    # row 0 = header, row 1 = pre-existing blank data row
    while len(table.rows) - 1 < len(rows):
        table.add_row()
    for i, ref in enumerate(rows):
        row = table.rows[i + 1]
        row.cells[0].text = ref.get("reference_number", "") or ""
        row.cells[1].text = ref.get("title", "") or ""
    # clear any unused pre-existing rows beyond the supplied refs
    for i in range(len(rows), len(table.rows) - 1):
        row = table.rows[i + 1]
        row.cells[0].text = ""
        row.cells[1].text = ""


def _add_table_column(table, header_text: str) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tbl = table._tbl
    grid = tbl.tblGrid
    grid_col = OxmlElement("w:gridCol")
    grid_col.set(qn("w:w"), "1600")
    grid.append(grid_col)
    for row in table.rows:
        tc = OxmlElement("w:tc")
        tc_pr = OxmlElement("w:tcPr")
        tc_w = OxmlElement("w:tcW")
        tc_w.set(qn("w:w"), "1600")
        tc_w.set(qn("w:type"), "dxa")
        tc_pr.append(tc_w)
        tc.append(tc_pr)
        p = OxmlElement("w:p")
        tc.append(p)
        row._tr.append(tc)
    table.rows[0].cells[-1].text = header_text


def _set_repeat_header(table) -> None:
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def fill_attendee_table(doc, attendees: list[dict]) -> None:
    """Table columns: [#, Attendee Name, Exam %, Certificate Awarded, Comments].

    A "Certificate Serial No." column is appended (the blank template does not
    have one) so every field the operator enters has its own cell. The table
    ships with 12 pre-built data rows; more are cloned via table.add_row() so
    attendance beyond 12 still renders correctly and continues across pages.
    """
    table = doc.tables[2]
    _add_table_column(table, "Certificate Serial No.")
    _set_repeat_header(table)

    rows = [a for a in attendees if (a.get("name") or "").strip() or (a.get("certificate_serial") or "").strip()]
    while len(table.rows) - 1 < len(rows):
        table.add_row()

    for i, att in enumerate(rows):
        row = table.rows[i + 1]
        row.cells[0].text = str(i + 1)
        row.cells[1].text = att.get("name", "") or ""
        exam_pct = att.get("exam_percentage")
        row.cells[2].text = f"{fmt_num(exam_pct)}%" if exam_pct not in (None, "") else ""
        row.cells[3].text = "Yes" if att.get("certificate_awarded") else "No"
        row.cells[4].text = att.get("certificate_serial", "") or ""
        row.cells[5].text = att.get("comments", "") or ""

    # blank out any unused pre-built rows beyond the supplied attendee count
    for i in range(len(rows), len(table.rows) - 1):
        row = table.rows[i + 1]
        for cell in row.cells[1:]:
            cell.text = ""


def insert_site_image(doc, image_bytes: Optional[bytes]) -> None:
    from docx.shared import Cm
    caption = None
    for p in doc.paragraphs:
        if _norm(p.text) == "Figure 1 site image":
            caption = p
            break
    if caption is None:
        return
    if not image_bytes:
        caption._p.getparent().remove(caption._p)
        return
    new_p = caption.insert_paragraph_before()
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    new_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = new_p.add_run()
    run.add_picture(io.BytesIO(image_bytes), width=Cm(14))
    caption.text = "Figure 1 site image"


def find_remaining_placeholders(doc) -> list[str]:
    found = []
    texts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                texts.append(cell.text)
    blob = "\n".join(texts)
    for token in KNOWN_PLACEHOLDER_TOKENS:
        if token in blob:
            found.append(token)
    return found


def build_training_course_report_docx(
    info_values: dict,
    doc_refs: list[dict],
    attendees: list[dict],
    narrative: dict,
    is_rtl: bool = False,
    site_image_bytes: Optional[bytes] = None,
) -> bytes:
    doc = load_template()

    fill_info_table(doc, info_values)
    fill_doc_ref_table(doc, doc_refs)
    fill_attendee_table(doc, attendees)
    insert_site_image(doc, site_image_bytes)

    for section_key in SECTION_HEADINGS:
        replace_section(doc, section_key, narrative.get(section_key), is_rtl=is_rtl)

    remaining = find_remaining_placeholders(doc)
    if remaining:
        raise RuntimeError(f"Template placeholders were not resolved: {remaining}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
