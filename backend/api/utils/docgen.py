"""
Document generation utilities for Innovation Engine exports.
Produces DOCX (via python-docx) and PDF (via WeasyPrint) files
with full Arabic RTL support.
"""
from __future__ import annotations
import io
import os
import re
from datetime import datetime, timezone
from urllib.parse import unquote, urlparse

# ── Content sanitiser ────────────────────────────────────────────────────────

_OFFER_PATTERNS = [
    r"If useful[,\s]+I can( next)? convert",
    r"If you('d| would) like[,\s]+I can",
    r"Let me know if you('d| would)? (like|want)",
    r"I can (also |next |now )?(generate|convert|create|produce|draft|write|turn|provide)",
    r"Would you like me to",
    r"Shall I (also|now|next|proceed|convert|create|generate|draft)",
    r"Do you want me to",
]
_OFFER_RE = re.compile(
    "(" + "|".join(_OFFER_PATTERNS) + ")",
    re.IGNORECASE,
)

_RUNTIME_NOISE_PATTERNS = [
    r"^Commercialisation analysis unavailable.*$",
    r"^External research did not finish within runtime budget.*$",
    r"^DRAFT\s+[—-]\s+EXTERNAL VERIFICATION INCOMPLETE.*$",
    r"^Runtime note:.*$",
]

_EXPORT_METADATA_PATTERNS = [
    r"^\s*Confidence\s*:\s*.*$",
    r"^\s*Official\s+Domain\s*:\s*.*$",
    r"^\s*No\s+verified\s+URL\s+available\s*$",
    r"^\s*Verified\s+via\s*:\s*.*$",
    r"^\s*Internal\s+verification\s+metadata\s*:\s*.*$",
]


def _strip_runtime_messages(content: str) -> str:
    """Remove runtime/configuration diagnostics from export payloads."""
    if not content:
        return content
    cleaned_lines = []
    for line in content.splitlines():
        text = line.strip()
        if text == "Chief X-Ray Innovation Scientist":
            continue
        if any(re.search(pat, text, flags=re.IGNORECASE) for pat in _RUNTIME_NOISE_PATTERNS):
            continue
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def _strip_offer_footer(content: str) -> str:
    """
    Remove any trailing 'offer to convert / follow-up action' paragraphs
    that GPT sometimes appends at the end of a document, before the first
    matched line and everything after it.
    """
    lines = content.splitlines()
    cutoff = None
    # Scan the LAST 30 lines only — the offer is always at the end
    tail_start = max(0, len(lines) - 30)
    for i, line in enumerate(lines[tail_start:], start=tail_start):
        stripped = line.strip()
        if stripped and _OFFER_RE.search(stripped):
            cutoff = i
            break
    if cutoff is not None:
        # Keep everything before the offer paragraph, strip trailing blank lines
        trimmed = lines[:cutoff]
        while trimmed and not trimmed[-1].strip():
            trimmed.pop()
        return "\n".join(trimmed)
    return content


def _strip_export_metadata(content: str) -> str:
    """Remove export-forbidden metadata lines from final paper payload."""
    if not content:
        return content
    cleaned_lines: list[str] = []
    for line in content.splitlines():
        text = line.strip()
        if any(re.search(pat, text, flags=re.IGNORECASE) for pat in _EXPORT_METADATA_PATTERNS):
            continue
        cleaned_lines.append(line)
    cleaned = "\n".join(cleaned_lines)
    return re.sub(r"\n{3,}", "\n\n", cleaned).strip()


def extract_paper_title(content: str, fallback: str = "Research Paper") -> str:
    """Extract a publication-style title from markdown content."""
    text = _strip_runtime_messages(_strip_offer_footer(content or "")).strip()
    if not text:
        return fallback

    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return fallback

    keyword_headers = {
        "abstract", "keywords", "introduction", "literature review", "related work",
        "novel contributions", "methodology", "mathematical models", "experimental setup",
        "results", "discussion", "limitations", "future work", "conclusion", "references",
    }

    for line in lines[:8]:
        plain = re.sub(r"^#{1,4}\s+", "", line).strip()
        if not plain:
            continue
        lowered = plain.lower().rstrip(":")
        if lowered in keyword_headers:
            continue
        if re.match(r"^(?:\d+\.|[IVX]+\.)\s+", plain):
            continue
        if plain.lower() in {"mohamed noaman", "x-ray academy research intelligence platform"}:
            continue
        if len(plain) < 8:
            continue
        return plain

    for line in lines:
        plain = re.sub(r"^#{1,4}\s+", "", line).strip()
        if plain and plain.lower() not in keyword_headers:
            return plain

    return fallback


def sanitize_export_filename(title: str, extension: str, max_length: int = 180) -> str:
    """Sanitize a paper title for safe use as a filename."""
    safe = (title or "Research Paper").strip()
    safe = safe.replace(":", " - ")
    safe = re.sub(r"[\\/*?<>|]", "", safe)
    safe = re.sub(r"\s+", " ", safe).strip()
    safe = safe.rstrip(". ")
    ext = extension if extension.startswith(".") else f".{extension.lstrip('.') }"
    reserved = {
        "CON", "PRN", "AUX", "NUL",
        *(f"COM{i}" for i in range(1, 10)),
        *(f"LPT{i}" for i in range(1, 10)),
    }
    if safe.upper() in reserved:
        safe = f"{safe} Paper"
    limit = max(1, max_length - len(ext))
    if len(safe) > limit:
        safe = safe[:limit].rstrip()
        safe = safe.rstrip(". ")
    return f"{safe}{ext}"


# ── Markdown parser ───────────────────────────────────────────────────────────

def _parse_md(md: str) -> list[dict]:
    """Parse markdown into [{level, title, body_lines}] blocks."""
    blocks: list[dict] = []
    cur: dict | None = None

    def flush():
        if cur is not None:
            cur["body"] = "\n".join(cur["body_lines"]).strip()
            blocks.append(cur)

    for line in md.splitlines():
        m1 = re.match(r'^(#{1,4})\s+(.*)', line)
        if m1:
            flush()
            level = len(m1.group(1))
            cur = {"level": level, "title": m1.group(2).strip(), "body_lines": []}
        else:
            if cur is None:
                cur = {"level": 0, "title": "", "body_lines": []}
            cur["body_lines"].append(line)
    flush()
    return blocks


# ── Arabic helpers ────────────────────────────────────────────────────────────

def _reshape_ar(text: str) -> str:
    """Reshape + apply BiDi algorithm so Arabic renders correctly."""
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        reshaped = arabic_reshaper.reshape(text)
        return get_display(reshaped)
    except Exception:
        return text


def _is_arabic(text: str) -> bool:
    """Heuristic: does this string contain Arabic characters?"""
    return bool(re.search(r'[\u0600-\u06FF]', text))


# ── DOCX helpers ──────────────────────────────────────────────────────────────

def _set_rtl_para(para, is_rtl: bool = True):
    """Apply RTL direction + right-alignment to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    pPr = para._p.get_or_add_pPr()
    # bidi element
    bidi_el = pPr.find(qn('w:bidi'))
    if bidi_el is None:
        bidi_el = OxmlElement('w:bidi')
        pPr.append(bidi_el)
    bidi_el.set(qn('w:val'), '1' if is_rtl else '0')

    if is_rtl:
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_heading_docx(doc, text: str, level: int, is_rtl: bool, font_name: str):
    from docx.shared import Pt, RGBColor
    heading_styles = {1: ('Heading 1', 16), 2: ('Heading 2', 13), 3: ('Heading 3', 11)}
    style_name, pt = heading_styles.get(level, ('Heading 2', 13))
    try:
        para = doc.add_heading(level=min(level, 4))
        run = para.runs[0] if para.runs else para.add_run()
        run.text = _reshape_ar(text) if is_rtl else text
        run.font.name = font_name
        run.font.size = Pt(pt)
    except Exception:
        para = doc.add_paragraph(text)
    if is_rtl:
        _set_rtl_para(para)
    return para


def _add_body_docx(doc, body: str, is_rtl: bool, font_name: str):
    """Add body text: handles bullet lists, blank lines, and plain paragraphs."""
    from docx.shared import Pt, Inches

    def _append_omath_to_paragraph(para, expression: str):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn

        expr = (expression or "").strip()
        if not expr:
            return
        omath_para = OxmlElement('m:oMathPara')
        omath = OxmlElement('m:oMath')
        mr = OxmlElement('m:r')
        mt = OxmlElement('m:t')
        mt.text = expr
        mr.append(mt)
        omath.append(mr)
        omath_para.append(omath)
        para._p.append(omath_para)

    def _add_table_from_md(table_lines: list[str]):
        rows: list[list[str]] = []
        for raw in table_lines:
            if '|' not in raw:
                continue
            cells = [c.strip() for c in raw.strip().strip('|').split('|')]
            if cells and all(set(c) <= set('-: ') for c in cells):
                continue
            rows.append(cells)
        if not rows:
            return
        ncols = max(len(r) for r in rows)
        padded = [r + [''] * (ncols - len(r)) for r in rows]
        table = doc.add_table(rows=len(padded), cols=ncols)
        table.style = 'Table Grid'
        for r_idx, row in enumerate(padded):
            for c_idx, cell_text in enumerate(row):
                text = re.sub(r'\*\*(.+?)\*\*', r'\1', cell_text)
                text = re.sub(r'\*(.+?)\*', r'\1', text)
                text = re.sub(r'`(.+?)`', r'\1', text)
                text = _reshape_ar(text) if is_rtl else text
                cell = table.cell(r_idx, c_idx)
                cell.text = text
                for para in cell.paragraphs:
                    if para.runs:
                        para.runs[0].font.name = font_name
                        para.runs[0].font.size = Pt(10)
                    if is_rtl:
                        _set_rtl_para(para)

    inline_eq_re = re.compile(r'\$([^$\n]{1,300})\$|\\\((.+?)\\\)')
    image_re = re.compile(r'^!\[(.*?)\]\((.+?)\)$')

    def _resolve_image_path(path_value: str) -> str:
        raw = (path_value or "").strip().strip('"').strip("'")
        if raw.lower().startswith("file://"):
            parsed = urlparse(raw)
            raw = unquote(parsed.path.lstrip('/'))
        return raw
    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        i += 1
        if not line:
            continue

        # Fenced code block. Figures are now rendered as real images
        # (see api/utils/figure_renderer.py) and referenced via markdown
        # image syntax, handled above — a fence should never contain a
        # diagram at this point. Guard anyway: a stray ```mermaid block must
        # never leak into the export as literal per-line paragraph text
        # (the previous behavior when no fence handling existed at all).
        if line.startswith('```'):
            lang = line[3:].strip().lower()
            fence_lines: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                fence_lines.append(lines[i])
                i += 1
            if i < len(lines):
                i += 1  # consume closing fence
            if lang == 'mermaid':
                continue
            code_text = "\n".join(fence_lines).strip()
            if code_text:
                para = doc.add_paragraph()
                run = para.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            continue

        img_match = image_re.match(line)
        if img_match:
            alt_text = img_match.group(1).strip() or "Figure"
            img_path = _resolve_image_path(img_match.group(2))
            if os.path.isfile(img_path):
                try:
                    doc.add_picture(img_path, width=Inches(6.2))
                    cap = doc.add_paragraph(alt_text)
                    if cap.runs:
                        cap.runs[0].font.name = font_name
                        cap.runs[0].font.size = Pt(10)
                    if is_rtl:
                        _set_rtl_para(cap)
                    continue
                except Exception:
                    pass

        # Markdown table -> native Word table
        if line.startswith('|') and '|' in line:
            table_lines = [line]
            while i < len(lines):
                nxt = lines[i].strip()
                if not nxt.startswith('|'):
                    break
                table_lines.append(nxt)
                i += 1
            _add_table_from_md(table_lines)
            continue

        # Block LaTeX equations -> native Word equation objects
        block_expr = None
        block_match = re.match(r'^\$\$(.+)\$\$$', line)
        if block_match:
            block_expr = block_match.group(1).strip()
        else:
            block_match = re.match(r'^\\\[(.+)\\\]$', line)
            if block_match:
                block_expr = block_match.group(1).strip()
        if block_expr:
            para = doc.add_paragraph()
            _append_omath_to_paragraph(para, block_expr)
            continue
        # Bullet
        if line.startswith(('- ', '* ', '• ')):
            text = line[2:].strip()
            if is_rtl:
                text = _reshape_ar(text)
            try:
                para = doc.add_paragraph(text, style='List Bullet')
            except Exception:
                para = doc.add_paragraph(f"• {text}")
            para.runs[0].font.name = font_name if para.runs else None
            if is_rtl:
                _set_rtl_para(para)
        # Numbered list
        elif re.match(r'^\d+\.', line):
            text = re.sub(r'^\d+\.\s*', '', line)
            if is_rtl:
                text = _reshape_ar(text)
            try:
                para = doc.add_paragraph(text, style='List Number')
            except Exception:
                para = doc.add_paragraph(f"  {text}")
            if para.runs:
                para.runs[0].font.name = font_name
            if is_rtl:
                _set_rtl_para(para)
        else:
            # Bold/inline markdown strip
            text = re.sub(r'\*\*(.+?)\*\*', r'\1', line)
            text = re.sub(r'\*(.+?)\*', r'\1', text)
            text = re.sub(r'`(.+?)`', r'\1', text)
            para = doc.add_paragraph()
            pos = 0
            has_inline_eq = False
            for m in inline_eq_re.finditer(text):
                has_inline_eq = True
                start, end = m.span()
                if start > pos:
                    plain = text[pos:start]
                    if plain:
                        run = para.add_run(_reshape_ar(plain) if is_rtl else plain)
                        run.font.name = font_name
                        run.font.size = Pt(11)
                expr = (m.group(1) or m.group(2) or "").strip()
                _append_omath_to_paragraph(para, expr)
                pos = end
            if has_inline_eq and pos < len(text):
                tail = text[pos:]
                if tail:
                    run = para.add_run(_reshape_ar(tail) if is_rtl else tail)
                    run.font.name = font_name
                    run.font.size = Pt(11)
            if not has_inline_eq:
                plain = _reshape_ar(text) if is_rtl else text
                run = para.add_run(plain)
                run.font.name = font_name
                run.font.size = Pt(11)
            if is_rtl:
                _set_rtl_para(para)


# ── DOCX builder ──────────────────────────────────────────────────────────────

def _enforce_mandatory_sections_sync(content: str) -> str:
    """
    Synchronous guard used by export functions.
    If the content is missing sections 13-17 (the mandatory pipeline tail),
    append a clear notice so the exported document never silently omits them.
    This is a last-resort fallback — correctly-generated reports always have
    these sections already present from the pipeline.
    """
    from api.utils.mandatory_sections import _MANDATORY_HEADING_RE
    if _MANDATORY_HEADING_RE.search(content):
        return content  # mandatory sections already present

    notice = (
        "\n\n---\n\n"
        "## 13. References\n\n"
        "[1] This report was generated before the mandatory sections pipeline was active. "
        "Re-generate to obtain a fully populated §13–§17.\n\n"
        "## 14. Related Patents\n\n"
        "_Re-generate this report to populate Related Patents._\n\n"
        "## 15. Standards\n\n"
        "_Re-generate this report to populate Standards._\n\n"
        "## 16. Knowledge Base Sources\n\n"
        "_Re-generate this report to populate Knowledge Base Sources._\n\n"
        "## 17. Commercialisation\n\n"
        "_Re-generate this report to populate Commercialisation analysis._\n"
    )
    return content.rstrip() + notice


def _publication_mode_label(mode: str) -> str:
    labels = {
        "paper_ieee": "IEEE Research Paper",
        "paper_elsevier": "Elsevier Journal Article",
        "literature_review": "Systematic Literature Review",
        "research_ideas": "Novel Research Ideas",
        "research_gaps": "Research Gap Analysis",
        "experiment_plan": "Experiment Plan",
        "patent_analysis": "Patent & Prior-Art Analysis",
        "patent_claims": "Patent Claims Drafting",
        "technical_report": "Technical Report",
        "security_algorithm": "Security Algorithm Design",
        "research": "Research Paper",
    }
    return labels.get(mode, mode.replace("_", " ").title())


def build_docx(
    record,
    content: str,
    lang: str = "en",
    layout: str = "sequential",
    ar_content: str | None = None,
) -> bytes:
    """
    Build a .docx file for the given innovation record.
    lang: 'en' | 'ar' | 'bilingual'
    layout (bilingual only): 'sequential' | 'sidebyside'
    ar_content: pre-translated Arabic markdown (required when lang != 'en')
    """
    content = _strip_offer_footer(content)
    content = _strip_runtime_messages(content)
    content = _strip_export_metadata(content)
    content = _enforce_mandatory_sections_sync(content)
    paper_title = extract_paper_title(content, fallback=getattr(record, "topic", "Research Paper"))
    if ar_content:
        ar_content = _strip_offer_footer(ar_content)
        ar_content = _strip_runtime_messages(ar_content)
        ar_content = _strip_export_metadata(ar_content)
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2.5)

    # Font choices
    en_font = "Arial"
    ar_font = "Arial"  # Arial has Arabic glyphs in Word

    now_str = datetime.now(timezone.utc).strftime("%d %B %Y")
    doc_id = record.id[:8].upper()
    mode_label = _publication_mode_label(record.mode)

    is_ar = (lang == "ar")
    is_bi = (lang == "bilingual")
    is_publication_paper = str(record.mode).startswith("paper_")

    # ── Cover page ────────────────────────────────────────────────────────────
    if is_publication_paper:
        cover_items = [
            ("X-Ray Academy", 28, True, True),
            ("Research Paper", 14, True, False),
            ("Mohamed Noaman", 12, True, False),
            ("X-Ray Academy Research Intelligence Platform", 10, False, False),
            ("", 10, False, False),
            (paper_title, 18, True, False),
            ("", 10, False, False),
            (f"Manuscript Type: {mode_label}", 12, False, False),
            (f"Technology Domain: {record.domain_label}", 12, False, False),
            (f"Document ID: {doc_id}", 12, False, False),
            (f"Date: {now_str}", 12, False, False),
            ("", 10, False, False),
            ("Publication-quality scientific manuscript prepared for journal-style review.", 11, True, False),
        ]
    else:
        cover_items = [
            ("X-Ray Academy", 28, True, True),
            ("Innovation Engine", 14, True, False),
            ("Research by Mohamed Noaman", 11, False, False),
            ("", 12, False, False),
            (paper_title, 18, True, False),
            ("", 12, False, False),
            (f"Technology Domain: {record.domain_label}", 12, False, False),
            (f"Generation Mode: {mode_label}", 12, False, False),
            (f"Document ID: {doc_id}", 12, False, False),
            (f"Date: {now_str}", 12, False, False),
            ("", 12, False, False),
            ("CONFIDENTIAL — For internal use only.", 11, True, False),
            ("This document contains proprietary innovation concepts and may constitute patentable subject matter.", 10, False, False),
        ]

    for text, size, bold, big in cover_items:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = en_font
        run.font.size = Pt(size)
        run.font.bold = bold

    doc.add_page_break()

    # ── Table of contents (manual) ────────────────────────────────────────────
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run("Table of Contents")
    toc_run.font.name = en_font
    toc_run.font.size = Pt(16)
    toc_run.font.bold = True

    blocks = _parse_md(content)
    toc_headings = [(b["level"], b["title"]) for b in blocks if b["level"] in (1, 2) and b["title"]]
    for lvl, title in toc_headings:
        indent = "    " if lvl > 1 else ""
        p = doc.add_paragraph(f"{indent}{title}")
        if p.runs:
            p.runs[0].font.name = en_font
            p.runs[0].font.size = Pt(10)

    doc.add_page_break()

    # ── Content ───────────────────────────────────────────────────────────────
    if lang == "en":
        for block in blocks:
            if block["level"] > 0 and block["title"]:
                _add_heading_docx(doc, block["title"], block["level"], False, en_font)
            if block["body"]:
                _add_body_docx(doc, block["body"], False, en_font)
            doc.add_paragraph()

    elif lang == "ar":
        ar_blocks = _parse_md(ar_content or content)
        for block in ar_blocks:
            if block["level"] > 0 and block["title"]:
                _add_heading_docx(doc, block["title"], block["level"], True, ar_font)
            if block["body"]:
                _add_body_docx(doc, block["body"], True, ar_font)
            doc.add_paragraph()

    elif lang == "bilingual":
        ar_blocks = _parse_md(ar_content or content)
        en_blocks = blocks

        if layout == "sequential":
            # Arabic section first, then English
            doc.add_heading("Arabic / العربية", level=1)
            for block in ar_blocks:
                if block["level"] > 0 and block["title"]:
                    _add_heading_docx(doc, block["title"], block["level"], True, ar_font)
                if block["body"]:
                    _add_body_docx(doc, block["body"], True, ar_font)
                doc.add_paragraph()

            doc.add_page_break()
            doc.add_heading("English", level=1)
            for block in en_blocks:
                if block["level"] > 0 and block["title"]:
                    _add_heading_docx(doc, block["title"], block["level"], False, en_font)
                if block["body"]:
                    _add_body_docx(doc, block["body"], False, en_font)
                doc.add_paragraph()

        else:  # sidebyside — use a 2-column table per section pair
            for en_b, ar_b in zip(en_blocks, ar_blocks + [{}] * len(en_blocks)):
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                left_cell = table.cell(0, 0)
                right_cell = table.cell(0, 1)

                # EN left
                en_title = en_b.get("title", "")
                en_body = en_b.get("body", "")
                if en_title:
                    p = left_cell.add_paragraph(en_title)
                    if p.runs:
                        p.runs[0].font.bold = True
                        p.runs[0].font.name = en_font
                if en_body:
                    left_cell.add_paragraph(en_body[:1000])

                # AR right
                ar_title = _reshape_ar(ar_b.get("title", "")) if ar_b else ""
                ar_body_raw = ar_b.get("body", "") if ar_b else ""
                ar_body = _reshape_ar(ar_body_raw[:1000])
                if ar_title:
                    p = right_cell.add_paragraph(ar_title)
                    if p.runs:
                        p.runs[0].font.bold = True
                        p.runs[0].font.name = ar_font
                    _set_rtl_para(p)
                if ar_body:
                    p = right_cell.add_paragraph(ar_body)
                    if p.runs:
                        p.runs[0].font.name = ar_font
                    _set_rtl_para(p)

                doc.add_paragraph()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── PDF builder (ReportLab / Platypus) ───────────────────────────────────────

# Colour palette
_NAVY   = (0.039, 0.086, 0.157)   # #0a1628
_BLUE   = (0.118, 0.227, 0.373)   # #1e3a5f
_GREY   = (0.4, 0.4, 0.4)
_WHITE  = (1, 1, 1)
_LIGHT  = (0.96, 0.97, 0.98)


def _rl_styles(is_rtl: bool = False):
    """Return a dict of ReportLab ParagraphStyles."""
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER, TA_JUSTIFY
    from reportlab.lib.units import mm
    from reportlab.lib import colors

    align = TA_RIGHT if is_rtl else TA_JUSTIFY
    heading_align = TA_RIGHT if is_rtl else TA_LEFT

    base = getSampleStyleSheet()

    def _s(name, parent="Normal", **kw) -> ParagraphStyle:
        return ParagraphStyle(name, parent=base[parent], **kw)

    return {
        "h1": _s("H1", "Heading1",
                 fontSize=16, textColor=colors.Color(*_NAVY),
                 spaceAfter=6, spaceBefore=18, alignment=heading_align,
                 fontName="Helvetica-Bold", leading=20),
        "h2": _s("H2", "Heading2",
                 fontSize=13, textColor=colors.Color(*_BLUE),
                 spaceAfter=4, spaceBefore=14, alignment=heading_align,
                 fontName="Helvetica-Bold", leading=17),
        "h3": _s("H3", "Heading3",
                 fontSize=11, textColor=colors.Color(*_BLUE),
                 spaceAfter=3, spaceBefore=10, alignment=heading_align,
                 fontName="Helvetica-Bold", leading=14),
        "body": _s("Body", fontSize=10, leading=15,
                   spaceAfter=6, alignment=align,
                   fontName="Helvetica"),
        "bullet": _s("Bullet", fontSize=10, leading=14,
                     spaceAfter=3, leftIndent=14, firstLineIndent=0,
                     bulletIndent=4, alignment=align, fontName="Helvetica"),
        "code": _s("Code", fontSize=8, leading=12,
                   fontName="Courier", backColor=colors.Color(0.94, 0.96, 0.98),
                   spaceAfter=6, leftIndent=10, rightIndent=10),
        "center": _s("Center", fontSize=10, leading=14, alignment=TA_CENTER,
                     fontName="Helvetica"),
        "cover_title": _s("CoverTitle", fontSize=22, leading=28,
                           textColor=colors.Color(*_NAVY), alignment=TA_CENTER,
                           fontName="Helvetica-Bold", spaceAfter=12),
        "cover_sub": _s("CoverSub", fontSize=12, leading=16,
                         textColor=colors.Color(*_GREY), alignment=TA_CENTER,
                         fontName="Helvetica", spaceAfter=8),
        "cover_meta": _s("CoverMeta", fontSize=10, leading=14,
                          textColor=colors.Color(*_GREY), alignment=TA_CENTER,
                          fontName="Helvetica", spaceAfter=4),
        "toc_entry": _s("TocEntry", fontSize=10, leading=14,
                         spaceAfter=2, fontName="Helvetica"),
        "toc_sub": _s("TocSub", fontSize=9, leading=12,
                       spaceAfter=2, leftIndent=20, fontName="Helvetica",
                       textColor=colors.Color(*_GREY)),
    }


# ── LaTeX math constants (used by _sanitize_latex_for_reportlab) ─────────────
# Use [$$] notation to avoid backslash-dollar edge cases in some text editors.
_LATEX_INLINE_RE        = re.compile(r'\\\((.+?)\\\)',       re.DOTALL)
_LATEX_BLOCK_RE         = re.compile(r'\\\[(.+?)\\\]',       re.DOTALL)
_LATEX_DOLLAR_BLOCK_RE  = re.compile(r'[$]{2}(.+?)[$]{2}',  re.DOTALL)
_LATEX_DOLLAR_INLINE_RE = re.compile(r'[$]([^$\n]{1,300})[$]')


def _sanitize_latex_for_reportlab(text: str) -> str:
    """
    Remove LaTeX math delimiters and make math content safe for ReportLab's
    XML-like paragraph parser.

    ReportLab's Paragraph() treats text as simplified XML, so literal '<' and '>'
    inside LaTeX expressions (e.g. C<C_min) corrupt the parse tree.
    This must run BEFORE any <font>/<b> markup is injected by _clean_inline.
    """
    def _plain(expr: str) -> str:
        expr = expr.replace('&', '&amp;')
        expr = expr.replace('<', '&lt;')
        expr = expr.replace('>', '&gt;')
        # Strip LaTeX decorators — keep base symbol, drop the markup
        expr = re.sub(r'\\hat\{([^}]+)\}',              lambda m: m.group(1) + '\u0302', expr)
        expr = re.sub(r'\\vec\{([^}]+)\}',              lambda m: m.group(1),            expr)
        expr = re.sub(r'\\bar\{([^}]+)\}',              lambda m: m.group(1) + '\u0304', expr)
        expr = re.sub(r'\\tilde\{([^}]+)\}',            lambda m: m.group(1) + '\u0303', expr)
        expr = re.sub(r'\\mathbf\{([^}]+)\}',           lambda m: m.group(1),            expr)
        expr = re.sub(r'\\text\{([^}]+)\}',             lambda m: m.group(1),            expr)
        expr = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', lambda m: '(' + m.group(1) + ')/(' + m.group(2) + ')', expr)
        expr = re.sub(r'\\([a-zA-Z]+)\{([^}]*)\}',     lambda m: m.group(2),            expr)
        expr = re.sub(r'\\[a-zA-Z]+',                   '',                              expr)
        expr = expr.replace('{', '').replace('}', '')
        return expr.strip()

    text = _LATEX_INLINE_RE.sub(        lambda m: _plain(m.group(1)),             text)
    text = _LATEX_BLOCK_RE.sub(         lambda m: '[' + _plain(m.group(1)) + ']', text)
    text = _LATEX_DOLLAR_BLOCK_RE.sub(  lambda m: '[' + _plain(m.group(1)) + ']', text)
    text = _LATEX_DOLLAR_INLINE_RE.sub( lambda m: _plain(m.group(1)),             text)
    # Final sweep: escape any bare < or > that weren't inside a LaTeX delimiter.
    # Safe here because _clean_inline has not yet injected <font>/<b> markup.
    text = text.replace('<', '&lt;').replace('>', '&gt;')
    return text


def _clean_inline(text: str) -> str:
    """Sanitize LaTeX math, then apply markdown inline \u2192 ReportLab XML markup."""
    # LaTeX must be sanitized BEFORE <font> tags are injected; LaTeX can contain raw '<' '>'
    text = _sanitize_latex_for_reportlab(text)
    text = re.sub(r'\*\*(.+?)\*\*', r'<b>\1</b>', text)
    text = re.sub(r'\*(.+?)\*',    r'<i>\1</i>', text)
    text = re.sub(r'`(.+?)`',      r'\1', text)
    text = re.sub(r'\[ESTABLISHED\]', '<font color="#059669"><b>[ESTABLISHED]</b></font>', text)
    text = re.sub(r'\[HYPOTHESIS\]',  '<font color="#D97706"><b>[HYPOTHESIS]</b></font>', text)
    text = re.sub(r'\[SPECULATIVE\]', '<font color="#7C3AED"><b>[SPECULATIVE]</b></font>', text)
    return text


def _md_to_flowables(md: str, styles: dict, is_rtl: bool = False) -> list:
    """Convert markdown to a list of ReportLab Platypus flowables."""
    from reportlab.platypus import (
        Paragraph, Spacer, Table, TableStyle, HRFlowable, Image
    )
    from reportlab.lib.units import mm
    from reportlab.lib import colors

    flowables = []
    lines = md.splitlines()
    in_table: bool = False
    table_rows: list[list[str]] = []
    pending_list: list[str] = []
    list_style: str = ""
    image_re = re.compile(r'^!\[(.*?)\]\((.+?)\)$')

    def _resolve_image_path(path_value: str) -> str:
        raw = (path_value or "").strip().strip('"').strip("'")
        if raw.lower().startswith("file://"):
            parsed = urlparse(raw)
            raw = unquote(parsed.path.lstrip('/'))
        return raw

    def flush_list():
        nonlocal pending_list, list_style
        for item in pending_list:
            text = _reshape_ar(item) if is_rtl else item
            text = _clean_inline(text)
            bullet = "• " if list_style == "ul" else ""
            flowables.append(Paragraph(f"{bullet}{text}", styles["bullet"]))
        pending_list = []
        list_style = ""

    def flush_table():
        nonlocal in_table, table_rows
        if not table_rows:
            in_table = False
            return
        # Filter out separator rows
        data_rows = []
        for row in table_rows:
            cells = [c.strip() for c in row]
            if not all(set(c) <= set('-|: ') for c in cells):
                data_rows.append(cells)

        if not data_rows:
            table_rows = []
            in_table = False
            return

        # Pad all rows to same column count
        ncols = max(len(r) for r in data_rows)
        padded = [r + [''] * (ncols - len(r)) for r in data_rows]

        tbl = Table(padded, repeatRows=1)
        tbl.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(*_NAVY)),
            ('TEXTCOLOR',  (0, 0), (-1, 0), colors.white),
            ('FONTNAME',   (0, 0), (-1, 0), 'Helvetica-Bold'),
            ('FONTSIZE',   (0, 0), (-1, -1), 9),
            ('GRID',       (0, 0), (-1, -1), 0.5, colors.Color(0.8, 0.8, 0.8)),
            ('ROWBACKGROUNDS', (0, 1), (-1, -1),
             [colors.white, colors.Color(*_LIGHT)]),
            ('VALIGN',     (0, 0), (-1, -1), 'TOP'),
            ('LEFTPADDING', (0, 0), (-1, -1), 4),
            ('RIGHTPADDING', (0, 0), (-1, -1), 4),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        flowables.append(Spacer(1, 4))
        flowables.append(tbl)
        flowables.append(Spacer(1, 8))
        table_rows = []
        in_table = False

    in_code = False
    code_lines: list[str] = []
    code_lang = ""

    for line in lines:
        # Fenced code block. Figures are rendered as real images (see
        # api/utils/figure_renderer.py) and referenced via markdown image
        # syntax, handled below — a fence must never leak into the PDF as
        # literal per-line paragraph text (the previous behavior when no
        # fence handling existed at all).
        stripped_line = line.strip()
        if stripped_line.startswith('```'):
            if in_code:
                in_code = False
                code_text = "\n".join(code_lines).strip()
                code_lines = []
                if code_lang != 'mermaid' and code_text:
                    flowables.append(Paragraph(_clean_inline(code_text).replace("\n", "<br/>"), styles["code"]))
                code_lang = ""
            else:
                flush_list()
                in_code = True
                code_lang = stripped_line[3:].strip().lower()
            continue
        if in_code:
            code_lines.append(line)
            continue

        # Table row detection
        if '|' in line and line.strip().startswith('|'):
            flush_list()
            cells = [c for c in line.strip().split('|') if c != '']
            table_rows.append(cells)
            in_table = True
            continue
        elif in_table:
            flush_table()

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            flush_list()
            lvl = len(m.group(1))
            title = m.group(2).strip()
            if is_rtl:
                title = _reshape_ar(title)
            style_key = {1: "h1", 2: "h2", 3: "h3"}.get(lvl, "h3")
            flowables.append(Spacer(1, 4))
            flowables.append(Paragraph(title, styles[style_key]))
            if lvl <= 2:
                flowables.append(HRFlowable(width="100%", thickness=0.5,
                                             color=colors.Color(0.8, 0.85, 0.9)))
            continue

        img_match = image_re.match(line.strip())
        if img_match:
            flush_list()
            alt_text = img_match.group(1).strip() or "Figure"
            img_path = _resolve_image_path(img_match.group(2))
            if os.path.isfile(img_path):
                try:
                    img = Image(img_path)
                    img._restrictSize(160 * mm, 110 * mm)
                    flowables.append(Spacer(1, 4))
                    flowables.append(img)
                    flowables.append(Paragraph(_clean_inline(alt_text), styles["center"]))
                    flowables.append(Spacer(1, 6))
                    continue
                except Exception:
                    pass

        # Bullets
        if line.strip().startswith(('- ', '* ', '• ')):
            if list_style and list_style != "ul":
                flush_list()
            list_style = "ul"
            pending_list.append(line.strip()[2:].strip())
            continue

        # Numbered list
        nm = re.match(r'^\d+\.\s+(.*)', line.strip())
        if nm:
            if list_style and list_style != "ol":
                flush_list()
            list_style = "ol"
            pending_list.append(nm.group(1))
            continue

        # Blank line
        if not line.strip():
            flush_list()
            flowables.append(Spacer(1, 6))
            continue

        # Normal paragraph
        flush_list()
        text = line.strip()
        if is_rtl:
            text = _reshape_ar(text)
        text = _clean_inline(text)
        flowables.append(Paragraph(text, styles["body"]))

    flush_list()
    flush_table()
    return flowables


def _page_header_footer(canvas, doc):
    """Draw header and footer on each page."""
    from reportlab.lib.units import mm
    from reportlab.lib import colors

    canvas.saveState()
    w, h = canvas.bookmarkPage, None
    page_w = doc.pagesize[0]
    page_h = doc.pagesize[1]

    # Header line
    canvas.setStrokeColor(colors.Color(*_NAVY))
    canvas.setLineWidth(0.5)
    canvas.line(20*mm, page_h - 15*mm, page_w - 20*mm, page_h - 15*mm)
    canvas.setFont("Helvetica", 8)
    canvas.setFillColor(colors.Color(*_GREY))
    canvas.drawString(20*mm, page_h - 13*mm, "X-Ray Academy · Innovation Engine")
    canvas.drawRightString(page_w - 20*mm, page_h - 13*mm, "CONFIDENTIAL")

    # Footer line
    canvas.line(20*mm, 15*mm, page_w - 20*mm, 15*mm)
    canvas.drawCentredString(page_w / 2, 10*mm, f"Page {doc.page}")

    canvas.restoreState()


def build_pdf(
    record,
    content: str,
    lang: str = "en",
    layout: str = "sequential",
    ar_content: str | None = None,
) -> bytes:
    """Build a PDF via ReportLab Platypus. Returns raw PDF bytes."""
    content = _strip_offer_footer(content)
    content = _strip_export_metadata(content)
    content = _enforce_mandatory_sections_sync(content)
    paper_title = extract_paper_title(content, fallback=getattr(record, "topic", "Research Paper"))
    if ar_content:
        ar_content = _strip_offer_footer(ar_content)
        ar_content = _strip_export_metadata(ar_content)
    import io as _io
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        PageBreak, HRFlowable
    )
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors

    buf = _io.BytesIO()
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=25*mm, rightMargin=20*mm,
        topMargin=22*mm, bottomMargin=20*mm,
    )

    now_str = datetime.now(timezone.utc).strftime("%d %B %Y")
    doc_id = record.id[:8].upper()
    mode_label = _publication_mode_label(record.mode)
    is_publication_paper = str(record.mode).startswith("paper_")

    is_ar = (lang == "ar")
    styles_en = _rl_styles(is_rtl=False)
    styles_ar = _rl_styles(is_rtl=True)

    story = []

    # ── Cover page ────────────────────────────────────────────────────────────
    story.append(Spacer(1, 40*mm))
    if is_publication_paper:
        story.append(Paragraph("X-Ray Academy", styles_en["cover_title"]))
        story.append(Paragraph("Research Paper", styles_en["cover_sub"]))
        story.append(Paragraph("Mohamed Noaman", styles_en["cover_sub"]))
        story.append(Paragraph("X-Ray Academy Research Intelligence Platform", styles_en["cover_meta"]))
    else:
        story.append(Paragraph("X-Ray Academy", styles_en["cover_title"]))
        story.append(Paragraph("Innovation Engine",
                                styles_en["cover_sub"]))
        story.append(Paragraph("Research by Mohamed Noaman", styles_en["cover_meta"]))
    story.append(Spacer(1, 10*mm))
    story.append(HRFlowable(width="80%", thickness=2,
                              color=colors.Color(*_NAVY), hAlign='CENTER'))
    story.append(Spacer(1, 8*mm))
    story.append(Paragraph(paper_title, styles_en["cover_title"]))
    story.append(Spacer(1, 8*mm))
    story.append(HRFlowable(width="60%", thickness=0.5,
                              color=colors.Color(*_GREY), hAlign='CENTER'))
    story.append(Spacer(1, 6*mm))
    for label, value in [
        ("Domain",        record.domain_label),
        ("Mode",          mode_label),
        ("Document ID",   doc_id),
        ("Date",          now_str),
        ("Word Count",    f"{record.word_count:,}"),
        ("KB Chunks",     str(record.kb_chunks_used)),
    ]:
        story.append(Paragraph(f"<b>{label}:</b>  {value}", styles_en["cover_meta"]))
    story.append(Spacer(1, 10*mm))

    # Confidentiality box
    conf_data = [["CONFIDENTIAL — For internal use only.\n"
                  "This document contains proprietary innovation concepts and may constitute "
                  "patentable subject matter. Do not distribute without authorisation."]]
    conf_tbl = Table(conf_data, colWidths=[160*mm])
    conf_tbl.setStyle(TableStyle([
        ('BOX',         (0, 0), (-1, -1), 0.5, colors.Color(0.7, 0.7, 0.7)),
        ('BACKGROUND',  (0, 0), (-1, -1), colors.Color(0.97, 0.97, 0.97)),
        ('FONTNAME',    (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE',    (0, 0), (-1, -1), 9),
        ('TEXTCOLOR',   (0, 0), (-1, -1), colors.Color(*_GREY)),
        ('ALIGN',       (0, 0), (-1, -1), 'CENTER'),
        ('LEFTPADDING', (0, 0), (-1, -1), 10),
        ('RIGHTPADDING',(0, 0), (-1, -1), 10),
        ('TOPPADDING',  (0, 0), (-1, -1), 8),
        ('BOTTOMPADDING',(0,0), (-1, -1), 8),
    ]))
    story.append(conf_tbl)
    story.append(PageBreak())

    # ── Table of contents ─────────────────────────────────────────────────────
    blocks = _parse_md(content)
    toc_headings = [(b["level"], b["title"]) for b in blocks if b["level"] in (1, 2) and b["title"]]

    story.append(Paragraph("Table of Contents", styles_en["h1"]))
    story.append(HRFlowable(width="100%", thickness=1, color=colors.Color(*_NAVY)))
    story.append(Spacer(1, 6))
    for lvl, title in toc_headings:
        style = styles_en["toc_sub"] if lvl > 1 else styles_en["toc_entry"]
        story.append(Paragraph(title, style))
    story.append(PageBreak())

    # ── Content ───────────────────────────────────────────────────────────────
    if lang == "en":
        story.extend(_md_to_flowables(content, styles_en, is_rtl=False))

    elif lang == "ar":
        story.extend(_md_to_flowables(ar_content or content, styles_ar, is_rtl=True))

    else:  # bilingual
        ar = ar_content or content
        if layout == "sequential":
            story.append(Paragraph(_reshape_ar("العربية"), styles_ar["h1"]))
            story.append(HRFlowable(width="100%", thickness=1,
                                     color=colors.Color(*_NAVY)))
            story.extend(_md_to_flowables(ar, styles_ar, is_rtl=True))
            story.append(PageBreak())
            story.append(Paragraph("English", styles_en["h1"]))
            story.append(HRFlowable(width="100%", thickness=1,
                                     color=colors.Color(*_NAVY)))
            story.extend(_md_to_flowables(content, styles_en, is_rtl=False))
        else:  # side-by-side — single multi-row 2-col table (Paragraph cells only)
            ar_blocks = _parse_md(ar)
            en_blocks = blocks
            col_w = 82*mm

            def _strip_md(t: str) -> str:
                t = re.sub(r'\*\*(.+?)\*\*', r'\1', t)
                t = re.sub(r'\*(.+?)\*',    r'\1', t)
                t = re.sub(r'`(.+?)`',      r'\1', t)
                return t

            table_data = [
                [Paragraph("English", styles_en["h2"]),
                 Paragraph(_reshape_ar("العربية"), styles_ar["h2"])],
            ]
            for en_b, ar_b in zip(en_blocks, ar_blocks + [{}] * len(en_blocks)):
                en_title = en_b.get("title", "")
                en_body  = _strip_md(en_b.get("body", "")[:600])
                ar_title = _reshape_ar(ar_b.get("title", "")) if ar_b else ""
                ar_body  = _reshape_ar(_strip_md((ar_b.get("body", "") if ar_b else "")[:600]))

                en_parts = ([f"<b>{en_title}</b>"] if en_title else []) + ([en_body.replace('\n','<br/>')] if en_body else [])
                ar_parts = ([f"<b>{ar_title}</b>"] if ar_title else []) + ([ar_body.replace('\n','<br/>')] if ar_body else [])
                table_data.append([
                    Paragraph("<br/>".join(en_parts) or " ", styles_en["body"]),
                    Paragraph("<br/>".join(ar_parts) or " ", styles_ar["body"]),
                ])

            main_tbl = Table(table_data, colWidths=[col_w, col_w])
            main_tbl.setStyle(TableStyle([
                ('BACKGROUND',   (0, 0), (-1, 0), colors.Color(*_NAVY)),
                ('TEXTCOLOR',    (0, 0), (-1, 0), colors.white),
                ('FONTNAME',     (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('VALIGN',       (0, 0), (-1, -1), 'TOP'),
                ('GRID',         (0, 0), (-1, -1), 0.3, colors.Color(0.85, 0.85, 0.85)),
                ('ROWBACKGROUNDS',(0, 1), (-1, -1), [colors.white, colors.Color(*_LIGHT)]),
                ('LEFTPADDING',  (0, 0), (-1, -1), 5),
                ('RIGHTPADDING', (0, 0), (-1, -1), 5),
                ('TOPPADDING',   (0, 0), (-1, -1), 4),
                ('BOTTOMPADDING',(0, 0), (-1, -1), 4),
            ]))
            story.append(main_tbl)

    doc.build(story, onFirstPage=_page_header_footer, onLaterPages=_page_header_footer)
    return buf.getvalue()
